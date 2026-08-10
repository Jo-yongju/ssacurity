from __future__ import annotations

import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "STM32_모터제어_완전입문서.docx"
ASSET_DIR = ROOT / "tmp" / "motor_control_book_assets"

KOREAN_FONT = "Malgun Gothic"
CODE_FONT = "Consolas"
FONT_REGULAR = Path(r"C:\Windows\Fonts\malgun.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\malgunbd.ttf")

NAVY = "163A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
SKY = "E8F1F8"
PALE_BLUE = "F2F7FB"
INK = "20262E"
MUTED = "5F6B76"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D7DEE5"
GREEN = "287D5A"
PALE_GREEN = "EAF6EF"
GOLD = "8A6500"
PALE_GOLD = "FFF7DD"
RED = "9B1C1C"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"

# compact_reference_guide preset, with a named Korean-legibility font override.
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Column widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(
                cell,
                CELL_MARGIN_TOP_BOTTOM,
                CELL_MARGIN_START_END,
                CELL_MARGIN_TOP_BOTTOM,
                CELL_MARGIN_START_END,
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color: str, size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), color)


def set_keep_with_next(paragraph, enabled=True) -> None:
    paragraph.paragraph_format.keep_with_next = enabled


def set_run_font(run, name=KOREAN_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_text_style(cell, *, bold=False, color=INK, size=9.2, align=None):
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), KOREAN_FONT)
    r_fonts.set(qn("w:hAnsi"), KOREAN_FONT)
    r_fonts.set(qn("w:eastAsia"), KOREAN_FONT)
    r_pr.extend([r_fonts, color, underline])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, field_text, fld_char_end])
    set_run_font(run, size=8.5, color=MUTED)


def add_bullet(doc: Document, text: str, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def _new_decimal_numbering_id(doc: Document) -> int:
    """Create a fresh decimal list so each procedural sequence restarts at 1."""
    numbering = doc.part.numbering_part.element
    decimal_abstract_id = None
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        level = abstract_num.find(qn("w:lvl"))
        if level is None or level.get(qn("w:ilvl")) != "0":
            continue
        num_fmt = level.find(qn("w:numFmt"))
        if num_fmt is not None and num_fmt.get(qn("w:val")) == "decimal":
            decimal_abstract_id = abstract_num.get(qn("w:abstractNumId"))
            break
    if decimal_abstract_id is None:
        raise RuntimeError("A decimal numbering definition was not found.")

    existing_ids = [
        int(num.get(qn("w:numId")))
        for num in numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) is not None
    ]
    num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), decimal_abstract_id)
    num.append(abstract_num_id)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_numbered_list(doc: Document, items: list[str]):
    num_id = _new_decimal_numbering_id(doc)
    paragraphs = []
    for text in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_element = OxmlElement("w:numId")
        num_id_element.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num_id_element])
        p.add_run(text)
        paragraphs.append(p)
    return paragraphs


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        p.add_run(text[len(bold_lead) :])
    else:
        p.add_run(text)
    return p


def add_callout(doc: Document, label: str, text: str, kind="note"):
    palette = {
        "note": (SKY, BLUE),
        "tip": (PALE_GREEN, GREEN),
        "warning": (PALE_GOLD, GOLD),
        "danger": (PALE_RED, RED),
    }
    fill, accent = palette[kind]
    p = doc.add_paragraph(style="Callout")
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, accent)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10.2, color=accent, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.2, color=INK)
    return p


def add_code_block(doc: Document, source: str, code: str, caption: str | None = None):
    if caption:
        p_caption = doc.add_paragraph(style="Code Caption")
        r = p_caption.add_run(caption)
        set_run_font(r, size=8.5, color=MUTED, italic=True)
    p = doc.add_paragraph(style="Code Block")
    set_paragraph_shading(p, "F6F8FA")
    set_paragraph_left_border(p, "A9B7C5", size=10)
    lines = code.strip("\n").splitlines()
    for index, line in enumerate(lines):
        run = p.add_run(line.rstrip())
        set_run_font(run, name=CODE_FONT, size=8.2, color="1F2933")
        if index != len(lines) - 1:
            run.add_break()
    if source:
        p_source = doc.add_paragraph(style="Source Path")
        r = p_source.add_run(f"현재 코드 위치: {source}")
        set_run_font(r, name=CODE_FONT, size=7.8, color=MUTED)
    return p


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
    *,
    header_fill=SKY,
    font_size=9.0,
    alignments: Sequence[WD_ALIGN_PARAGRAPH] | None = None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, header_fill)
        set_cell_text_style(
            cell,
            bold=True,
            color=NAVY,
            size=font_size,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        for col_index, value in enumerate(row_values):
            cell = row.cells[col_index]
            cell.text = str(value)
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            alignment = (
                alignments[col_index]
                if alignments is not None
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_cell_text_style(
                cell,
                bold=False,
                color=INK,
                size=font_size,
                align=alignment,
            )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(4)
    return table


def add_figure(doc: Document, path: Path, caption: str, alt_text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", alt_text)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)
    return p


def chapter(doc: Document, number: str, title: str, subtitle: str | None = None):
    p_kicker = doc.add_paragraph()
    p_kicker.paragraph_format.page_break_before = True
    p_kicker.paragraph_format.space_after = Pt(2)
    r = p_kicker.add_run(f"CHAPTER {number}")
    set_run_font(r, size=9.5, color=GOLD, bold=True)
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(title)
    if subtitle:
        sub = doc.add_paragraph(style="Chapter Subtitle")
        sub.add_run(subtitle)


def part_page(doc: Document, number: str, title: str, promise: str):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(120)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"PART {number}")
    set_run_font(r, size=11, color=GOLD, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(14)
    r = p2.add_run(title)
    set_run_font(r, size=27, color=NAVY, bold=True)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.left_indent = Inches(0.65)
    p3.paragraph_format.right_indent = Inches(0.65)
    r = p3.add_run(promise)
    set_run_font(r, size=12, color=MUTED)


def heading(doc: Document, level: int, title: str, *, page_break_before: bool = False):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break_before
    p.add_run(title)
    return p


def question_box(doc: Document, question: str, hint: str | None = None):
    text = question if not hint else f"{question}\n힌트: {hint}"
    return add_callout(doc, "생각해 보기", text, "note")


def draw_centered_text(draw, box, text, font, fill, max_lines=4):
    x0, y0, x1, y1 = box
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if draw.textlength(test, font=font) <= (x1 - x0 - 32):
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    lines = lines[:max_lines]
    line_h = font.size + 10
    total_h = len(lines) * line_h
    y = y0 + ((y1 - y0) - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x0 + ((x1 - x0) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h


def draw_box(draw, box, text, fill, outline, font, text_fill=INK, radius=22):
    draw.rounded_rectangle(box, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=4)
    draw_centered_text(draw, box, text, font, f"#{text_fill}")


def draw_arrow(draw, start, end, color=BLUE, width=8):
    x1, y1 = start
    x2, y2 = end
    draw.line((x1, y1, x2, y2), fill=f"#{color}", width=width)
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 22
    spread = 0.6
    points = [
        (x2, y2),
        (
            x2 - length * math.cos(angle - spread),
            y2 - length * math.sin(angle - spread),
        ),
        (
            x2 - length * math.cos(angle + spread),
            y2 - length * math.sin(angle + spread),
        ),
    ]
    draw.polygon(points, fill=f"#{color}")


def make_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    regular = ImageFont.truetype(str(FONT_REGULAR), 34)
    small = ImageFont.truetype(str(FONT_REGULAR), 27)
    bold = ImageFont.truetype(str(FONT_BOLD), 36)
    tiny = ImageFont.truetype(str(FONT_REGULAR), 23)

    outputs: dict[str, Path] = {}

    # 1. Whole-system data flow.
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 55), "명령이 모터 움직임으로 바뀌는 전체 흐름", font=bold, fill=f"#{NAVY}")
    boxes = [
        ((60, 220, 300, 400), "PC / Jetson\n명령", SKY, BLUE),
        ((360, 220, 620, 400), "CommTask\n1 ms", PALE_GREEN, GREEN),
        ((690, 220, 970, 400), "ControlTask\n10 ms", PALE_GOLD, GOLD),
        ((1040, 220, 1370, 400), "BTS7960\nDC 모터", PALE_RED, RED),
    ]
    for box, text, fill, outline in boxes:
        draw_box(draw, box, text, fill, outline, regular)
    for x1, x2 in [(300, 360), (620, 690), (970, 1040)]:
        draw_arrow(draw, (x1 + 8, 325), (x2 - 8, 325))
    draw_box(draw, (80, 600, 370, 790), "TIM2 초음파\n60 ms 측정", SKY, BLUE, small)
    draw_box(draw, (470, 600, 760, 790), "SafetyTask\n10 ms 판단", PALE_RED, RED, small)
    draw_box(draw, (870, 600, 1160, 790), "TIM4 엔코더\n속도 측정", PALE_GREEN, GREEN, small)
    draw_arrow(draw, (370, 695), (470, 695), color=RED)
    draw_arrow(draw, (760, 640), (825, 410), color=RED)
    draw_arrow(draw, (1200, 430), (1165, 640), color=BLUE)
    draw_arrow(draw, (1015, 600), (880, 410), color=GREEN)
    draw.text((90, 520), "거리 → 안전 요청 → 출력 허가", font=small, fill=f"#{RED}")
    draw.text((915, 520), "속도 피드백", font=small, fill=f"#{GREEN}")
    path = ASSET_DIR / "01_system_flow.png"
    img.save(path, dpi=(180, 180))
    outputs["system_flow"] = path

    # 2. Wiring map.
    img = Image.new("RGB", (1600, 1250), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 50), "현재 프로젝트의 모터·엔코더 배선 지도", font=bold, fill=f"#{NAVY}")
    draw.rounded_rectangle((70, 180, 510, 1170), radius=28, fill=f"#{SKY}", outline=f"#{BLUE}", width=5)
    draw.text((125, 220), "STM32F429I-DISC1", font=regular, fill=f"#{NAVY}")
    draw.rounded_rectangle((980, 180, 1510, 650), radius=28, fill=f"#{PALE_GOLD}", outline=f"#{GOLD}", width=5)
    draw.text((1090, 220), "BTS7960 모듈", font=regular, fill=f"#{NAVY}")
    labels = [
        ("PA0 / TIM5_CH1", "RPWM", 350, BLUE),
        ("PA3 / TIM5_CH4", "LPWM", 440, BLUE),
        ("PE2 / GPIO", "R_EN", 530, GREEN),
        ("PE3 / GPIO", "L_EN", 620, GREEN),
    ]
    for left, right, y, color in labels:
        draw.text((160, y - 18), left, font=small, fill=f"#{INK}")
        draw.text((1240, y - 18), right, font=small, fill=f"#{INK}")
        draw_arrow(draw, (510, y), (1200, y), color=color, width=6)
    draw.rounded_rectangle((980, 720, 1510, 920), radius=28, fill=f"#{PALE_GREEN}", outline=f"#{GREEN}", width=5)
    draw.text((1060, 750), "Quadrature Encoder", font=regular, fill=f"#{NAVY}")
    draw.text((1215, 820), "A / B", font=small, fill=f"#{INK}")
    draw.text((160, 745), "PB6 / TIM4_CH1", font=small, fill=f"#{INK}")
    draw.text((160, 835), "PB7 / TIM4_CH2", font=small, fill=f"#{INK}")
    draw_arrow(draw, (980, 785), (510, 785), color=GREEN, width=6)
    draw_arrow(draw, (980, 875), (510, 875), color=GREEN, width=6)
    draw.text((1035, 675), "모터 전원과 STM32 GND는 공통 기준", font=tiny, fill=f"#{RED}")
    draw.rounded_rectangle((980, 960, 1510, 1160), radius=28, fill="#F2F4F6", outline=f"#{MUTED}", width=5)
    draw.text((1070, 985), "Ultrasonic sensor", font=regular, fill=f"#{NAVY}")
    draw.text((1240, 1055), "TRIG", font=small, fill=f"#{INK}")
    draw.text((1240, 1115), "ECHO", font=small, fill=f"#{INK}")
    draw.text((160, 1010), "PA5 / GPIO", font=small, fill=f"#{INK}")
    draw.text((160, 1090), "PB3 / TIM2_CH2", font=small, fill=f"#{INK}")
    draw_arrow(draw, (510, 1030), (1200, 1030), color=BLUE, width=6)
    draw_arrow(draw, (1200, 1110), (510, 1110), color=RED, width=6)
    path = ASSET_DIR / "02_wiring_map.png"
    img.save(path, dpi=(180, 180))
    outputs["wiring_map"] = path

    # 3. PWM timing.
    img = Image.new("RGB", (1600, 850), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "PWM은 빠른 ON/OFF의 비율로 세기를 만든다", font=bold, fill=f"#{NAVY}")
    x0, x1 = 160, 1470
    rows = [
        (190, 0.20, "20% Duty"),
        (390, 0.50, "50% Duty"),
        (590, 0.80, "80% Duty"),
    ]
    periods = 5
    for y, duty, label in rows:
        draw.text((70, y + 25), label, font=small, fill=f"#{INK}")
        period_w = (x1 - x0) / periods
        for i in range(periods):
            sx = x0 + i * period_w
            hi_end = sx + period_w * duty
            draw.line((sx, y + 90, sx, y), fill=f"#{BLUE}", width=5)
            draw.line((sx, y, hi_end, y), fill=f"#{BLUE}", width=5)
            draw.line((hi_end, y, hi_end, y + 90), fill=f"#{BLUE}", width=5)
            draw.line((hi_end, y + 90, sx + period_w, y + 90), fill=f"#{BLUE}", width=5)
        draw.text((1250, y + 105), f"평균 에너지 ≈ {int(duty * 100)}%", font=tiny, fill=f"#{MUTED}")
    path = ASSET_DIR / "03_pwm_duty.png"
    img.save(path, dpi=(180, 180))
    outputs["pwm"] = path

    # 4. Quadrature encoder.
    img = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "A와 B의 순서를 보면 회전 방향을 알 수 있다", font=bold, fill=f"#{NAVY}")
    start_x = 230
    step = 150
    for direction, y0, title in [
        ("forward", 200, "정방향 예: A가 B보다 먼저 변함"),
        ("reverse", 510, "역방향 예: B가 A보다 먼저 변함"),
    ]:
        draw.text((70, y0 - 40), title, font=small, fill=f"#{INK}")
        for channel, offset, color in [("A", 0, BLUE), ("B", 95, GREEN)]:
            y = y0 + offset
            draw.text((120, y + 20), channel, font=bold, fill=f"#{color}")
            sequence = [0, 1, 1, 0, 0, 1, 1, 0, 0]
            if (direction == "forward" and channel == "B") or (
                direction == "reverse" and channel == "A"
            ):
                sequence = [0, 0, 1, 1, 0, 0, 1, 1, 0]
            for i in range(len(sequence) - 1):
                x = start_x + i * step
                yv = y if sequence[i] else y + 60
                yv2 = y if sequence[i + 1] else y + 60
                draw.line((x, yv, x + step, yv), fill=f"#{color}", width=6)
                draw.line((x + step, yv, x + step, yv2), fill=f"#{color}", width=6)
    path = ASSET_DIR / "04_encoder_phase.png"
    img.save(path, dpi=(180, 180))
    outputs["encoder"] = path

    # 5. Closed loop.
    img = Image.new("RGB", (1600, 800), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "폐루프 속도 제어: 틀린 만큼 다시 고친다", font=bold, fill=f"#{NAVY}")
    draw_box(draw, (70, 270, 350, 470), "목표 속도\n0.150 m/s", SKY, BLUE, regular)
    draw_box(draw, (480, 270, 780, 470), "오차와 P 보정\nKp × error", PALE_GOLD, GOLD, regular)
    draw_box(draw, (910, 270, 1190, 470), "PWM + 모터", PALE_RED, RED, regular)
    draw_box(draw, (1270, 270, 1530, 470), "실제 속도", PALE_GREEN, GREEN, regular)
    draw_arrow(draw, (350, 370), (480, 370))
    draw_arrow(draw, (780, 370), (910, 370))
    draw_arrow(draw, (1190, 370), (1270, 370))
    draw.line((1400, 480, 1400, 650, 620, 650, 620, 485), fill=f"#{GREEN}", width=7)
    draw_arrow(draw, (620, 650), (620, 480), color=GREEN)
    draw.text((780, 675), "엔코더가 측정한 속도를 되돌려 줌", font=small, fill=f"#{GREEN}")
    path = ASSET_DIR / "05_closed_loop.png"
    img.save(path, dpi=(180, 180))
    outputs["closed_loop"] = path

    # 6. RTOS timeline.
    img = Image.new("RGB", (1600, 1080), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "태스크마다 다른 주기와 우선순위를 기억하자", font=bold, fill=f"#{NAVY}")
    x0, x1 = 250, 1500
    scale = (x1 - x0) / 120.0
    timelines = [
        ("SafetyTask · High", 180, 10, RED),
        ("ControlTask · AboveNormal", 360, 10, BLUE),
        ("CommTask · Normal", 540, 1, GREEN),
        ("UltrasonicTask · Normal", 720, 60, NAVY),
        ("Telemetry · Comm 내부", 900, 50, GOLD),
    ]
    for label, y, period, color in timelines:
        draw.text((70, y - 20), label, font=small, fill=f"#{INK}")
        draw.line((x0, y + 35, x1, y + 35), fill="#C9D2DA", width=4)
        for t in range(0, 121, period):
            x = x0 + t * scale
            draw.rectangle((x - 6, y, x + 6, y + 70), fill=f"#{color}")
        draw.text((1300, y + 80), f"주기 {period} ms", font=tiny, fill=f"#{color}")
    draw.text((250, 1035), "0 ms", font=tiny, fill=f"#{MUTED}")
    draw.text((1390, 1035), "120 ms", font=tiny, fill=f"#{MUTED}")
    path = ASSET_DIR / "06_rtos_timeline.png"
    img.save(path, dpi=(180, 180))
    outputs["rtos"] = path

    # 7. Ultrasonic safety gate.
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.text((70, 45), "초음파 값이 모터 허가 또는 정지 요청으로 바뀌는 길", font=bold, fill=f"#{NAVY}")
    safety_boxes = [
        ((70, 220, 350, 410), "TIM2 ECHO\npulse width", SKY, BLUE),
        ((440, 220, 730, 410), "UltrasonicTask\n중앙값 필터", PALE_GREEN, GREEN),
        ((820, 200, 1130, 430), "SafetyTask\n거리·상태 판단", PALE_GOLD, GOLD),
        ((1220, 220, 1530, 410), "ControlTask\n출력 허가/차단", PALE_RED, RED),
    ]
    for box, text, fill, outline in safety_boxes:
        draw_box(draw, box, text, fill, outline, regular)
    for x1, x2 in [(350, 440), (730, 820), (1130, 1220)]:
        draw_arrow(draw, (x1 + 8, 315), (x2 - 8, 315))
    bands = [
        ((90, 590, 430, 730), "0.65 m 이상\nCLEAR 복귀", PALE_GREEN, GREEN),
        ((470, 590, 810, 730), "0.60 m 미만\nCAUTION", SKY, BLUE),
        ((850, 590, 1190, 730), "0.20 m 미만\nSTOP", PALE_GOLD, GOLD),
        ((1230, 590, 1530, 730), "0.20 m 미만\nE-STOP 래치", PALE_RED, RED),
    ]
    for box, text, fill, outline in bands:
        draw_box(draw, box, text, fill, outline, small)
    draw.text((90, 790), "STOP 해제: 0.30 m 이상인 정상 샘플 3회", font=small, fill=f"#{MUTED}")
    draw.text((850, 790), "E-STOP 해제: 정지·정상 센서·안전 거리에서 reset", font=small, fill=f"#{RED}")
    path = ASSET_DIR / "07_safety_gate.png"
    img.save(path, dpi=(180, 180))
    outputs["safety"] = path

    return outputs


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = KOREAN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), KOREAN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = KOREAN_FONT
    h1._element.rPr.rFonts.set(qn("w:ascii"), KOREAN_FONT)
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = rgb(NAVY)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = KOREAN_FONT
    h2._element.rPr.rFonts.set(qn("w:ascii"), KOREAN_FONT)
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = KOREAN_FONT
    h3._element.rPr.rFonts.set(qn("w:ascii"), KOREAN_FONT)
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = rgb(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = KOREAN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), KOREAN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
        style.font.size = Pt(10.4)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    custom_styles = {
        "Chapter Subtitle": (11.5, MUTED, False, 0, 14, 1.2),
        "Callout": (10.2, INK, False, 6, 6, 1.18),
        "Code Block": (8.2, INK, False, 5, 5, 1.0),
        "Code Caption": (8.5, MUTED, False, 6, 2, 1.0),
        "Source Path": (7.8, MUTED, False, 1, 8, 1.0),
        "Caption": (8.8, MUTED, False, 4, 10, 1.0),
        "Small Text": (8.8, MUTED, False, 0, 4, 1.15),
    }
    for name, (size, color, bold, before, after, line) in custom_styles.items():
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = KOREAN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), KOREAN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = line
        style.paragraph_format.widow_control = True

    # Explicit compact list geometry.
    for style_name, left, first in [
        ("List Bullet", Inches(0.375), Inches(-0.188)),
        ("List Number", Inches(0.375), Inches(-0.188)),
        ("List Bullet 2", Inches(0.70), Inches(-0.188)),
    ]:
        style = doc.styles[style_name]
        style.paragraph_format.left_indent = left
        style.paragraph_format.first_line_indent = first

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("STM32 모터 제어 완전 입문서  |  ssacurity-stm32-drive")
    set_run_font(hr, size=8.2, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("PAGE ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_page_field(fp, "PAGE")
    fr2 = fp.add_run(" / ")
    set_run_font(fr2, size=8.5, color=MUTED)
    add_page_field(fp, "NUMPAGES")


def build_document() -> None:
    assets = make_diagrams()
    doc = Document()
    setup_styles(doc)

    # Cover: editorial_cover pattern, compact_reference_guide body preset.
    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(95)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("완전 쌩뉴비를 위한")
    set_run_font(r, size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("STM32 모터 제어")
    set_run_font(r, size=31, color=NAVY, bold=True)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.paragraph_format.space_after = Pt(24)
    r = title2.add_run("코드 한 권으로 끝내기")
    set_run_font(r, size=25, color=BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.left_indent = Inches(0.45)
    subtitle.paragraph_format.right_indent = Inches(0.45)
    subtitle.paragraph_format.space_after = Pt(80)
    r = subtitle.add_run(
        "GPIO부터 PWM·엔코더·PID·초음파 안전정지·FreeRTOS·UART까지\n"
        "현재 프로젝트 코드를 그대로 따라가며 배우는 실전 입문서"
    )
    set_run_font(r, size=13.5, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run("대상 프로젝트  ssacurity-stm32-drive")
    set_run_font(r, size=10, color=NAVY, bold=True)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta2.add_run("코드 기준  main/e8acfdc + 현재 안전·초음파 작업 트리 · 2026-07-28")
    set_run_font(r, size=9, color=MUTED)

    # Preface.
    heading(doc, 1, "이 책을 시작하기 전에", page_break_before=True)
    add_body(
        doc,
        "이 책은 C 언어와 전자회로를 처음 보는 사람도 현재 저장소의 모터 제어 코드를 "
        "읽을 수 있도록 만든 프로젝트 전용 교재다. 일반적인 STM32 이론을 길게 늘어놓기보다, "
        "지금 보이는 핀·타이머·함수·숫자가 왜 필요한지 한 층씩 쌓아 올린다.",
    )
    add_callout(
        doc,
        "이 책의 최종 목표",
        "코드를 외우는 것이 아니다. PC에서 들어온 명령이 어떤 함수와 하드웨어를 거쳐 "
        "모터의 회전이 되고, 엔코더 값이 어떻게 다시 PWM을 고치는지 자기 말로 설명하는 것이 목표다.",
        "tip",
    )
    heading(doc, 2, "먼저 약속할 안전 수칙")
    for item in [
        "첫 시험은 반드시 바퀴를 바닥에서 띄우고, 사람이 회전체에 손을 대지 못하게 고정한다.",
        "배선 변경은 모터 전원과 USB 전원을 모두 끈 상태에서 한다.",
        "STM32 GND와 BTS7960 논리 GND는 공통으로 연결하되, 모터 전류가 흐르는 굵은 전력 배선과 신호 배선을 분리한다.",
        "현재 PA0는 보드의 B1 USER 버튼과 연결되어 있다. PWM 동작 중 B1을 누르지 않는다.",
        "물리 E-stop은 UART STOP 명령으로 대체하지 않는다.",
        "처음에는 ±20~30%의 짧은 시간 제한 진단만 사용한다. 생산용 주행 제한은 실측 전까지 잠가 둔다.",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "매우 중요",
        "이 책은 코드와 정적 설정을 설명한다. 실제 모터 전원 전압, 모듈의 정품 여부, 배선 굵기, "
        "퓨즈, 냉각, 기계적 고정 상태는 저장소만으로 확인할 수 없다. 실제 구동 전 별도 점검이 필요하다.",
        "danger",
    )

    heading(doc, 2, "추천 학습 순서")
    add_table(
        doc,
        ["날짜", "읽을 장", "해볼 일", "완료 기준"],
        [
            ["1일차", "1~3장", "GPIO·타이머·PWM 용어 익히기", "20 kHz 계산을 직접 할 수 있다"],
            ["2일차", "4~5장", "main.c와 motor.c 따라 읽기", "양수/음수 PWM 경로를 설명한다"],
            ["3일차", "6장", "엔코더 속도 계산 손으로 풀기", "11 count가 약 0.138 m/s임을 계산한다"],
            ["4일차", "7~8장", "PID와 ControlTask 흐름 그리기", "목표·측정·오차·PWM 관계를 설명한다"],
            ["5일차", "9~11장", "UART 진단과 안전 시험", "자동 정지와 텔레메트리를 확인한다"],
            ["6~7일차", "12~13장", "고장 추적·작은 수정 연습", "한 번에 한 변수만 바꿔 비교한다"],
        ],
        [1050, 1500, 3270, 3540],
        font_size=8.8,
    )

    # Validation snapshot.
    heading(doc, 1, "현재 프로젝트 검증 스냅샷", page_break_before=True)
    add_body(
        doc,
        "교재는 실행 코드와 `.ioc`가 일치하는 현재 작업 트리를 우선한다. 오래된 설명 문서와 충돌할 때는 "
        "실제 초기화 코드, HAL 핸들, 새 안전·초음파 파일을 기준으로 판단했다.",
    )
    heading(doc, 2, "현재 자원 사용 현황")
    add_table(
        doc,
        ["자원", "현재 용도", "충돌/주의", "근거"],
        [
            ["PA0 / TIM5_CH1", "RPWM 정방향 20 kHz", "B1 USER 버튼과 공유", ".ioc, main.c, motor.c, UM1670"],
            ["PA3 / TIM5_CH4", "LPWM 역방향 20 kHz", "온보드 LCD B5와 공유", ".ioc, MSP, ST 보드 회로도"],
            ["PE2 / PE3", "R_EN / L_EN", "보드 기능 충돌 없음", ".ioc, main.c, UM1670"],
            ["PB6 / TIM4_CH1", "엔코더 A", "온보드 SDRAM SDNE1과 공유", ".ioc, encoder.c, UM1670"],
            ["PB7 / TIM4_CH2", "엔코더 B", "보드 기능 충돌 없음", ".ioc, encoder.c, UM1670"],
            ["PA5 / GPIO", "초음파 TRIG", "보드 표상 free I/O", ".ioc, main.c, UM1670"],
            ["PB3 / TIM2_CH2", "초음파 ECHO", "SWO와 공유; SB9 상태 확인", ".ioc, MSP, UM1670"],
            ["TIM2", "1 MHz ECHO input capture", "IRQ 우선순위 6", ".ioc, ultrasonic.c"],
            ["TIM6", "HAL 1 ms time base", "모터 타이머와 분리됨", "hal_timebase_tim.c"],
            ["USART1 + DMA2", "PC 진단/텔레메트리", "IRQ 우선순위 5", ".ioc, MSP, interrupt 파일"],
        ],
        [1900, 2380, 2550, 2530],
        font_size=8.4,
    )
    heading(doc, 2, "발견된 위험과 이 책의 판단", page_break_before=True)
    add_table(
        doc,
        ["위험도", "문제", "가능한 증상", "현재 학습 기준"],
        [
            ["높음", "PA0에 버튼이 물리 연결", "PWM 중 버튼 누르면 핀이 GND로 당겨짐", "B1 절대 누르지 않기"],
            ["높음", "PA3가 LCD B5와 공유", "LCD 오동작, 핀 기능 충돌", "LCD를 사용하지 않는 전제"],
            ["높음", "PB6가 SDRAM SDNE1과 공유", "SDRAM 사용 불가/오동작", "SDRAM을 사용하지 않는 전제"],
            ["높음", "초음파 ECHO 전압 미확인", "3.3 V 한계 초과 시 MCU 손상 가능", "분압/레벨시프터와 실측 확인"],
            ["중간", "PB3가 SWO와 공유", "trace 출력 불가 또는 신호 간섭", "SB9 실제 상태와 debug 설정 확인"],
            ["중간", "인수인계의 LPWM 정보가 오래됨", "PA6/TIM13으로 잘못 배선", "현재는 PA3/TIM5_CH4"],
            ["중간", "엔코더 입력 필터 0·NOPULL", "노이즈성 속도 점프", "실차에서 필터/배선 검증"],
            ["중간", "생산용 drive limit이 0", "비영점 CMD_DRIVE 거부", "진단 경로로만 시험"],
            ["중간", "안전·초음파 코드가 미커밋 작업 트리", "문서 기준이 곧 바뀔 수 있음", "현재 파일과 .ioc를 함께 보존"],
            ["낮음", "TIM5 설명 주석은 생성 영역", "CubeMX 재생성 때 주석 소실", "값은 .ioc로 보존"],
        ],
        [1050, 2600, 3050, 2660],
        font_size=8.4,
    )
    add_callout(
        doc,
        "문서 불일치",
        "`docs/handover_2026-07-28.md`에는 LPWM이 PA6/TIM13_CH1로 남아 있지만, "
        "현재 `.ioc`, `main.c`, `motor.c`, 작업 트리는 PA3/TIM5_CH4다. 이 책 전체는 후자를 사용한다. "
        "또한 안전·초음파 기능은 아직 커밋되지 않은 현재 작업 트리 기준이다.",
        "warning",
    )

    # Contents.
    heading(doc, 1, "차례와 코드 지도", page_break_before=True)
    add_table(
        doc,
        ["부", "장", "핵심 질문"],
        [
            ["1부 기초", "1. 큰 그림", "명령은 어떻게 모터 회전이 되는가?"],
            ["1부 기초", "2. 하드웨어", "STM32와 BTS7960은 어떤 선으로 연결되는가?"],
            ["1부 기초", "3. PWM", "20 kHz와 Duty는 어떻게 계산되는가?"],
            ["2부 코드", "4. 부팅", "main()에서 ControlTask까지 어떻게 도착하는가?"],
            ["2부 코드", "5. motor.c", "양수·음수 명령이 두 PWM으로 어떻게 갈리는가?"],
            ["2부 코드", "6. encoder.c", "카운트가 속도로 어떻게 바뀌는가?"],
            ["2부 코드", "7. pid.c", "오차가 PWM 보정값으로 어떻게 바뀌는가?"],
            ["2부 코드", "8. 제어·안전 게이트", "10 ms마다 출력 허가를 어떻게 판단하는가?"],
            ["2부 코드", "9. 통신/RTOS", "초음파·안전·진단 태스크는 어떻게 연결되는가?"],
            ["3부 실습", "10. 숫자 한 바퀴", "150 mm/s 명령을 손으로 추적할 수 있는가?"],
            ["3부 실습", "11. 안전 시험", "명령·텔레메트리·자동 정지를 어떻게 확인하는가?"],
            ["3부 실습", "12. 문제 해결", "모터가 안 돌 때 어디부터 볼 것인가?"],
            ["3부 실습", "13. 다음 수정", "어떻게 작게 바꾸고 안전하게 검증하는가?"],
        ],
        [1500, 2300, 5560],
        font_size=8.7,
    )
    heading(doc, 2, "가장 먼저 기억할 파일 12개")
    file_map = [
        ("ssacurity-stm32-drive.ioc", "핀·클록·타이머 설정의 원본"),
        ("Core/Src/main.c", "부팅, 클록, 주변장치 초기화, RTOS 시작"),
        ("Core/Src/stm32f4xx_hal_msp.c", "핀 Alternate Function과 주변장치 클록"),
        ("Drivers/BSP/Src/motor.c", "BTS7960에 실제 PWM을 쓰는 가장 아래층"),
        ("Drivers/BSP/Src/encoder.c", "TIM4 카운트를 거리·속도로 변환"),
        ("App/Src/pid.c", "목표와 측정의 차이를 보정값으로 변환"),
        ("App/Src/task_control.c", "10 ms 제어 상태기계와 안전 요청 적용"),
        ("Drivers/BSP/Src/ultrasonic.c", "TIM2 ECHO pulse 폭을 인터럽트로 측정"),
        ("App/Src/task_ultrasonic.c", "60 ms 측정과 3개 중앙값 필터"),
        ("App/Src/safety.c", "거리·센서 상태를 CLEAR/STOP/E-STOP으로 판단"),
        ("App/Src/task_safety.c", "10 ms마다 안전 판단을 모터와 ControlTask에 전달"),
        ("App/Src/app_freertos.c", "통신 명령을 제어 명령으로 연결하고 텔레메트리 송신"),
    ]
    add_table(
        doc,
        ["파일", "역할", "파일", "역할"],
        [
            [file_map[i][0], file_map[i][1], file_map[i + 6][0], file_map[i + 6][1]]
            for i in range(6)
        ],
        [2150, 2530, 2150, 2530],
        font_size=7.6,
    )

    part_page(
        doc,
        "1",
        "아무것도 몰라도 괜찮다",
        "먼저 전기 신호와 타이머의 언어를 익힌다. 이 부분을 이해하면 뒤의 C 코드가 갑자기 읽히기 시작한다.",
    )

    # Chapter 1.
    chapter(doc, "1", "명령에서 바퀴까지, 큰 그림", "전체 시스템을 한 장의 지도처럼 보기")
    add_figure(
        doc,
        assets["system_flow"],
        "그림 1-1. 통신 명령, 제어 태스크, 모터, 엔코더가 만드는 한 바퀴",
        "PC 또는 Jetson 명령이 CommTask와 ControlTask를 지나 BTS7960 모터로 전달되고 엔코더 피드백이 되돌아오는 흐름",
    )
    heading(doc, 2, "현재 모터 제어는 다섯 문장으로 요약된다")
    for sentence in [
        "1) 밖에서 목표를 받는다: 예를 들어 ‘앞으로 0.15 m/s로 가라’.",
        "2) 지금 속도를 잰다: 엔코더 펄스가 10 ms 동안 몇 번 들어왔는지 센다.",
        "3) 목표와 실제를 비교한다: 느리면 PWM을 조금 올리고, 빠르면 조금 내린다.",
        "4) 초음파 안전 요청이 허가할 때만 그 PWM을 낸다.",
        "5) 이 과정을 10 ms마다 반복한다.",
    ]:
        add_body(doc, sentence)
    add_callout(
        doc,
        "초보자 비유",
        "샤워 물 온도를 맞추는 것과 같다. 목표 온도와 손으로 느낀 온도를 비교하고, 차가우면 뜨거운 물을 더 연다. "
        "모터 제어에서는 손이 엔코더이고, 수도꼭지가 PWM이다.",
        "tip",
    )
    heading(doc, 2, "오픈 루프와 클로즈드 루프")
    add_table(
        doc,
        ["구분", "무엇을 믿는가", "현재 코드의 모드", "장단점"],
        [
            ["Open loop", "내가 준 PWM 값", "CONTROL_MODE_OPEN_LOOP", "단순하지만 배터리·부하에 따라 속도가 달라짐"],
            ["Closed loop", "엔코더가 잰 실제 속도", "CONTROL_MODE_SPEED_PID", "목표 속도를 따라가지만 보정·튜닝이 필요"],
            ["Disabled", "출력하지 않음", "CONTROL_MODE_DISABLED", "두 PWM=0, Enable=LOW로 안전 정지"],
        ],
        [1500, 2200, 2500, 3160],
        font_size=9.0,
    )
    heading(doc, 2, "레이어라는 생각")
    add_body(
        doc,
        "프로젝트는 한 파일이 모든 일을 하지 않는다. 하드웨어에 가까운 `motor.c`와 `encoder.c`, "
        "수학을 담당하는 `pid.c`, 정책을 담당하는 `task_control.c`, 외부 명령을 담당하는 "
        "`app_freertos.c`가 층을 이룬다. 아래층은 ‘어떻게’를, 위층은 ‘언제·왜’를 결정한다.",
    )
    question_box(
        doc,
        "모터가 느려졌을 때 PWM을 올리는 판단은 motor.c와 task_control.c 중 어디에 있어야 할까?",
        "motor.c는 요청받은 PWM을 정확히 출력하는 하드웨어 드라이버다.",
    )

    # Chapter 2.
    chapter(doc, "2", "하드웨어와 배선부터 이해하기", "STM32, BTS7960, DC 모터, 엔코더의 역할")
    add_figure(
        doc,
        assets["wiring_map"],
        "그림 2-1. 현재 코드가 기대하는 모터·엔코더·초음파 핀 연결",
        "STM32의 PA0 PA3 PE2 PE3가 BTS7960으로 가고 PB6 PB7은 엔코더, PA5 PB3은 초음파 TRIG ECHO에 연결되는 배선",
    )
    heading(doc, 2, "STM32가 모터를 직접 돌리지 않는 이유")
    add_body(
        doc,
        "STM32 핀은 논리 신호를 만드는 작은 스위치다. 모터가 요구하는 큰 전류를 공급할 수 없다. "
        "BTS7960 모듈은 STM32의 작은 3.3 V 신호를 받아 배터리의 큰 전력을 모터 쪽에서 스위칭한다. "
        "즉 STM32는 지휘관, BTS7960은 힘을 쓰는 작업자다.",
    )
    add_callout(
        doc,
        "전압 호환",
        "BTS7960 원 칩 데이터시트의 IN/INH 입력은 TTL/CMOS 호환이며 high 판정 상한이 약 2.15 V이므로 "
        "STM32의 3.3 V high로 직접 구동할 수 있다. 다만 시판 모듈의 VCC·점퍼·입력 회로는 제품마다 다를 수 있어 "
        "모듈 회로도 확인이 필요하다.",
        "note",
    )
    heading(doc, 2, "초음파 센서 두 선")
    add_table(
        doc,
        ["STM32", "센서 쪽", "방향", "현재 코드"],
        [
            ["PA5 / GPIO", "TRIG", "STM32 → 센서", "10 µs HIGH pulse"],
            ["PB3 / TIM2_CH2", "ECHO", "센서 → STM32", "상승·하강 edge capture"],
        ],
        [2200, 1900, 2100, 3160],
        font_size=9.0,
    )
    add_callout(
        doc,
        "ECHO 전압은 반드시 실측",
        "저장소에는 사용 센서의 정확한 모델·회로도가 없다. 5 V ECHO를 내는 모듈이면 STM32 입력 전에 "
        "분압기나 레벨시프터가 필요할 수 있다. 데이터시트와 오실로스코프로 3.3 V 호환을 확인하기 전에는 연결하지 않는다.",
        "danger",
    )
    add_callout(
        doc,
        "PB3와 SWO",
        "PB3는 보드의 SWO 신호와 공유된다. UM1670에 따르면 SB9가 ON이면 CN2 SWO와 PB3가 연결된다. "
        "ECHO로 쓰기 전 실제 SB9 상태와 IDE의 trace 사용 여부를 확인한다.",
        "warning",
    )
    heading(doc, 2, "네 개의 제어선")
    add_table(
        doc,
        ["STM32", "BTS7960", "현재 의미", "코드에서 조작"],
        [
            ["PA0 / TIM5_CH1", "RPWM", "정방향 세기", "CCR1"],
            ["PA3 / TIM5_CH4", "LPWM", "역방향 세기", "CCR4"],
            ["PE2", "R_EN", "정방향 쪽 활성화", "HAL_GPIO_WritePin"],
            ["PE3", "L_EN", "역방향 쪽 활성화", "HAL_GPIO_WritePin"],
        ],
        [2200, 1800, 2500, 2860],
        font_size=9.0,
    )
    heading(doc, 2, "방향 전환의 핵심 안전 규칙")
    add_body(
        doc,
        "정방향 PWM과 역방향 PWM을 동시에 켜면 원치 않는 브리지 상태가 생길 수 있다. "
        "`Motor_SetPercent()`는 새 방향을 선택하기 전에 두 채널의 CCR을 모두 0으로 만든다. "
        "그 뒤 양수면 CH1만, 음수면 CH4만 설정한다.",
    )
    add_code_block(
        doc,
        "Drivers/BSP/Src/motor.c · Motor_SetPercent()",
        """
SetBothPwmChannelsToZero();

if (requested_percent > 0.0f) {
    __HAL_TIM_SET_COMPARE(&htim5, TIM_CHANNEL_1, compare);
} else if (requested_percent < 0.0f) {
    __HAL_TIM_SET_COMPARE(&htim5, TIM_CHANNEL_4, compare);
}
""",
        "방향을 바꾸기 전 두 PWM을 먼저 끄는 현재 코드의 핵심",
    )
    heading(doc, 2, "보드 내장 기능과 공유되는 핀")
    add_callout(
        doc,
        "PA0",
        "보드 B1 USER 버튼과 연결된다. 버튼을 누르면 PA0가 GND로 연결되는 구성이다. "
        "PWM 출력 중 버튼을 누르지 말아야 한다.",
        "danger",
    )
    add_callout(
        doc,
        "PA3",
        "공식 보드 자료에서 LCD의 B5 색상선과 공유된다. 현재 펌웨어는 LCD를 실제로 초기화하지 않는 전제에서 LPWM으로 쓴다.",
        "warning",
    )
    add_callout(
        doc,
        "PB6",
        "온보드 SDRAM의 SDNE1과 기판상 공유된다. 현재 엔코더 입력으로 재배치했으므로 온보드 SDRAM과 동시에 사용할 수 없다.",
        "warning",
    )
    heading(doc, 2, "전원 배선 체크리스트")
    wiring_checklist = [
        "STM32 GND ↔ BTS7960 logic GND 공통",
        "모터 전원 +/− 극성 확인, 적절한 퓨즈 사용",
        "모터 전류 경로는 굵고 짧게, 신호선은 전력선과 떨어뜨리기",
        "모터 양단 또는 모듈 권장 위치의 노이즈 억제 부품 확인",
        "BTS7960 방열판과 케이블 허용 전류 확인",
        "엔코더 출력 전압이 STM32 3.3 V 입력 한계를 넘지 않는지 확인",
    ]
    for index, item in enumerate(wiring_checklist):
        paragraph = add_bullet(doc, item)
        paragraph.paragraph_format.keep_with_next = index < len(wiring_checklist) - 1

    # Chapter 3.
    chapter(doc, "3", "타이머와 PWM을 숫자로 이해하기", "20 kHz, ARR, CCR, Duty를 손으로 계산")
    add_figure(
        doc,
        assets["pwm"],
        "그림 3-1. 주파수는 빠르기, Duty는 한 주기 안에서 켜져 있는 비율",
        "20퍼센트 50퍼센트 80퍼센트 듀티의 PWM 파형 비교",
    )
    heading(doc, 2, "클록 트리에서 TIM5까지")
    add_body(
        doc,
        "시스템 클록은 HSI 16 MHz를 PLL로 올려 180 MHz가 된다. APB1 버스는 4분주되어 45 MHz지만, "
        "APB prescaler가 1이 아닐 때 STM32F4의 APB 타이머 클록은 버스 클록의 2배가 된다. "
        "따라서 TIM5 입력 클록은 90 MHz다.",
    )
    add_table(
        doc,
        ["단계", "설정", "결과"],
        [
            ["HSI", "16 MHz", "PLL 입력 원본"],
            ["PLL", "M=8, N=180, P=2", "SYSCLK 180 MHz"],
            ["APB1", "HCLK / 4", "PCLK1 45 MHz"],
            ["APB1 Timer", "prescaler ≠ 1이므로 ×2", "TIM5CLK 90 MHz"],
            ["TIM5", "PSC=8, ARR=499", "20 kHz"],
        ],
        [1800, 3900, 3660],
        font_size=9.0,
    )
    heading(doc, 2, "PWM 주파수 공식")
    add_callout(
        doc,
        "공식",
        "PWM 주파수 = Timer clock ÷ (PSC + 1) ÷ (ARR + 1)",
        "note",
    )
    add_body(
        doc,
        "현재 값은 90,000,000 ÷ 9 ÷ 500 = 20,000 Hz다. 즉 1초에 20,000번 PWM 주기가 반복된다. "
        "BTS7960 원 칩 데이터시트의 PWM 능력 25 kHz 이내에도 들어간다.",
    )
    heading(doc, 2, "Duty를 CCR 값으로 바꾸기")
    add_body(
        doc,
        "`PercentToCompare()`는 ARR+1을 한 주기의 count 수로 본다. ARR=499이므로 period_counts=500이다. "
        "30%라면 500×0.30=150, 22%라면 110을 CCR에 쓴다.",
    )
    add_table(
        doc,
        ["요청 Duty", "계산", "CCR", "의미"],
        [
            ["0%", "500 × 0.00", "0", "항상 LOW"],
            ["20%", "500 × 0.20", "100", "한 주기의 20% HIGH"],
            ["22%", "500 × 0.22", "110", "현재 최소 구동 feed-forward"],
            ["30%", "500 × 0.30", "150", "일반 진단 예"],
            ["40%", "500 × 0.40", "200", "현재 제어 안전 상한"],
            ["100%", "500 × 1.00", "500", "사실상 항상 HIGH"],
        ],
        [1700, 2600, 1500, 3560],
        font_size=9.0,
    )
    add_code_block(
        doc,
        "Drivers/BSP/Src/motor.c · PercentToCompare()",
        """
uint32_t period_counts = __HAL_TIM_GET_AUTORELOAD(timer) + 1U;
float compare = (magnitude_percent * (float)period_counts) / 100.0f;
return (uint32_t)(compare + 0.5f);
""",
        "반올림까지 포함한 Duty → CCR 변환",
    )
    question_box(
        doc,
        "PWM을 10 kHz로 낮추고 ARR=499를 유지하려면 PSC는 몇이어야 할까?",
        "90 MHz ÷ (PSC+1) ÷ 500 = 10 kHz",
    )

    part_page(
        doc,
        "2",
        "이제 실제 코드를 읽는다",
        "부팅부터 PWM 출력까지 파일을 건너다니며 따라간다. 각 함수가 무엇을 책임지고 무엇을 책임지지 않는지도 함께 본다.",
    )

    # Chapter 4.
    chapter(doc, "4", "main()에서 ControlTask까지", "전원이 켜진 뒤 초기화 순서 읽기")
    heading(doc, 2, "부팅 순서 한 줄 요약")
    add_code_block(
        doc,
        "Core/Src/main.c · main()",
        """
HAL_Init();
SystemClock_Config();
MX_GPIO_Init();
MX_DMA_Init();
MX_USART1_UART_Init();
MX_TIM2_Init();
MX_TIM3_Init();
MX_TIM4_Init();
MX_TIM5_Init();
CommService_Init(&huart1);
osKernelInitialize();
App_FreeRTOS_Init();
osKernelStart();
""",
        "현재 프로젝트의 핵심 초기화 순서",
    )
    heading(doc, 2, "왜 순서가 중요한가")
    for title, text in [
        ("HAL_Init()", "HAL 내부 상태와 1 ms time base를 준비한다. 이 프로젝트에서 HAL tick은 TIM6가 만든다."),
        ("SystemClock_Config()", "CPU와 버스, 타이머가 사용할 실제 주파수를 결정한다."),
        ("MX_GPIO_Init()", "Enable 핀을 먼저 LOW로 두어 부팅 중 모터가 켜지지 않게 한다."),
        ("MX_TIM2/4/5_Init()", "TIM2 초음파 입력 캡처, TIM4 엔코더, TIM5 PWM의 모드·주기·핀 AF를 준비한다. 아직 실제 측정·PWM은 시작하지 않는다."),
        ("CommService_Init()", "USART1 DMA 수신을 준비한다. 실패하면 Error_Handler로 간다."),
        ("App_FreeRTOS_Init()", "Control·Safety·Ultrasonic 모듈을 초기화하고 SafetyTask, ControlTask, UltrasonicTask, CommTask를 만든다."),
        ("osKernelStart()", "이후 실행 주도권이 FreeRTOS로 넘어간다."),
    ]:
        add_body(doc, f"{title} — {text}", bold_lead=title)
    heading(doc, 2, "타이머를 Init하는 것과 Start하는 것은 다르다")
    add_body(
        doc,
        "`MX_TIM5_Init()`는 TIM5의 레지스터 구성과 핀 Alternate Function을 준비한다. "
        "실제 카운터와 PWM 출력 시작은 나중에 `Motor_Init()`의 `HAL_TIM_PWM_Start()`가 한다. "
        "마찬가지로 TIM4 엔코더 카운터 시작은 `Encoder_Init()`에서, TIM2 입력 캡처 인터럽트 시작은 "
        "`Ultrasonic_Init()`에서 한다.",
    )
    add_callout(
        doc,
        "읽기 요령",
        "CubeMX 생성 함수 이름이 보이면 ‘설정 준비’, BSP의 Init 함수가 보이면 ‘실제 사용 시작’으로 먼저 나누어 생각하면 쉽다.",
        "tip",
    )
    heading(doc, 2, "CubeMX 재생성에 살아남는 구조")
    add_body(
        doc,
        "`CommService_Init()`과 `App_FreeRTOS_Init()` 호출은 `USER CODE` 영역에 있다. "
        "`App/`와 `Drivers/BSP/` 아래의 사용자 파일도 CubeMX 생성 대상이 아니다. "
        "반면 핀·타이머 값은 `.ioc`와 생성 코드가 함께 맞아야 한다.",
    )
    add_callout(
        doc,
        "주의",
        "`main.c`의 TIM5 주파수 설명 주석은 USER CODE 블록 밖에 있으므로 CubeMX 재생성 때 사라질 수 있다. "
        "실제 설정값은 `.ioc`의 Prescaler=8, Period=499가 기준이다.",
        "warning",
    )

    # Chapter 5.
    chapter(doc, "5", "motor.c 완전 해부", "요청 퍼센트를 안전한 두 채널 PWM으로 바꾸기")
    heading(doc, 2, "이 파일의 책임")
    add_body(
        doc,
        "`motor.c`는 속도 목표나 PID를 모른다. 오직 ‘몇 %로 어느 방향으로 출력하라’는 요청을 "
        "TIM5 CCR과 Enable 핀으로 바꾼다. 이런 분리가 있어야 상위 제어를 바꾸어도 하드웨어 드라이버를 재사용할 수 있다.",
    )
    heading(doc, 2, "상태 변수 다섯 개")
    add_table(
        doc,
        ["변수", "의미", "초기값", "왜 필요한가"],
        [
            ["motor_initialized", "PWM start 성공 여부", "false", "초기화 전 잘못된 출력을 막음"],
            ["motor_enabled", "Enable 핀 상태", "false", "비활성 상태에서 SetPercent를 무시"],
            ["motor_emergency_disabled", "비상 정지 latch", "false", "명시적 안전 해제 전 재출력을 막음"],
            ["motor_direction_inverted", "설치 방향 반전", "false", "배선을 바꾸지 않고 좌표계 보정"],
            ["motor_applied_percent", "논리적으로 적용된 명령", "0", "텔레메트리에 보고"],
        ],
        [2500, 2450, 1300, 3110],
        font_size=8.4,
    )
    heading(doc, 2, "Motor_Init()의 안전 순서")
    add_numbered_list(doc, [
        "R_EN과 L_EN을 LOW로 만든다.",
        "정·역방향 CCR을 0으로 만든다.",
        "TIM5 CH1 PWM을 시작한다.",
        "TIM5 CH4 PWM을 시작한다.",
        "둘 다 성공한 뒤에만 initialized=true로 기록한다.",
    ])
    add_code_block(
        doc,
        "Drivers/BSP/Src/motor.c · Motor_Init()",
        """
HAL_GPIO_WritePin(GPIOE, BTS_R_EN_Pin | BTS_L_EN_Pin, GPIO_PIN_RESET);
SetBothPwmChannelsToZero();

if (HAL_TIM_PWM_Start(&htim5, TIM_CHANNEL_1) != HAL_OK) return false;
if (HAL_TIM_PWM_Start(&htim5, TIM_CHANNEL_4) != HAL_OK) {
    HAL_TIM_PWM_Stop(&htim5, TIM_CHANNEL_1);
    return false;
}
""",
        "한 채널만 시작된 반쪽 상태도 정리하고 실패를 보고한다",
    )
    heading(doc, 2, "Motor_Enable()과 Motor_Disable()")
    add_body(
        doc,
        "`Motor_Enable()`은 두 Enable 핀을 HIGH로 만든다. `Motor_Disable()`은 더 강하다. "
        "두 CCR을 먼저 0으로 하고 두 Enable을 LOW로 내리며, 보고 값도 0으로 만든다. "
        "따라서 정지·fault·모드 전환에서 같은 안전 동작을 재사용할 수 있다.",
    )
    heading(doc, 2, "일반 정지와 비상 정지는 다르다")
    add_table(
        doc,
        ["함수", "무엇을 끄나", "다시 켜지는 조건"],
        [
            ["Motor_Disable()", "PWM=0, Enable=LOW", "다음 정상 Motor_Enable()"],
            ["Motor_EmergencyDisable()", "PWM=0, Enable=LOW, latch=true", "안전 reset 뒤 Clear 성공"],
            ["Motor_ClearEmergencyDisable()", "출력을 0으로 재확인하고 latch 해제", "motor_initialized=true"],
        ],
        [2900, 3300, 3160],
        font_size=8.8,
    )
    add_body(
        doc,
        "`motor_emergency_disabled`가 true면 `Motor_Enable()`과 `Motor_SetPercent()`가 출력을 거부한다. "
        "따라서 다른 태스크가 오래된 명령을 다시 써도 비상 정지 상태를 쉽게 깨지 못한다. "
        "공유 상태를 읽고 쓸 때는 PRIMASK로 매우 짧게 인터럽트를 막아 값과 핀 변경을 한 덩어리로 만든다.",
    )
    add_code_block(
        doc,
        "Drivers/BSP/Src/motor.c · emergency latch",
        """
motor_emergency_disabled = true;
SetBothPwmChannelsToZero();
HAL_GPIO_WritePin(BTS_R_EN_GPIO_Port,
                  BTS_R_EN_Pin | BTS_L_EN_Pin,
                  GPIO_PIN_RESET);
motor_enabled = false;
motor_applied_percent = 0.0f;
""",
        "비상 정지는 출력 0과 Enable LOW를 만든 뒤 래치를 남긴다",
    )
    heading(doc, 2, "Motor_SetPercent()를 문장으로 번역")
    add_numbered_list(doc, [
        "초기화 전·Enable OFF·비상 래치 중이면 두 PWM을 0으로 하고 끝낸다.",
        "요청을 -100~+100%로 제한한다.",
        "설치 방향 반전 옵션이 켜졌으면 부호를 뒤집는다.",
        "방향 전환 전 두 PWM을 0으로 한다.",
        "양수면 CH1, 음수면 CH4에 절댓값 Duty를 쓴다.",
        "외부에 보고할 applied percent를 갱신한다.",
    ])

    # Chapter 6.
    chapter(doc, "6", "encoder.c 완전 해부", "16비트 카운터에서 m/s 속도까지")
    add_figure(
        doc,
        assets["encoder"],
        "그림 6-1. Quadrature A/B의 위상 순서가 방향을 만든다",
        "A와 B 두 채널의 위상 차이로 정방향과 역방향을 구분하는 파형",
    )
    heading(doc, 2, "TIM4 Encoder Mode TI12")
    add_body(
        doc,
        "TIM4는 PB6(CH1)과 PB7(CH2)의 두 신호를 하드웨어로 비교해 카운터를 올리거나 내린다. "
        "CPU가 모든 edge마다 인터럽트를 처리하지 않아도 되어 효율적이다. 현재 ARR=65535인 16비트 카운터다.",
    )
    heading(doc, 2, "왜 65535에서 0으로 넘어가도 괜찮을까")
    add_body(
        doc,
        "카운터가 65535에서 한 칸 증가하면 0이 된다. 단순히 `0-65535`를 계산하면 큰 음수처럼 보인다. "
        "현재 코드는 unsigned 16비트 뺄셈 뒤 int16_t로 해석한다. 제어 주기 사이 변화량이 ±32767 이내라는 전제에서 "
        "wrap 전후의 작은 변화량을 그대로 복원할 수 있다.",
    )
    add_code_block(
        doc,
        "Drivers/BSP/Src/encoder.c · Encoder_Update()",
        """
current_counter = (uint16_t)__HAL_TIM_GET_COUNTER(&htim4);
wrapped_delta = (int16_t)(uint16_t)(current_counter - previous_counter);
previous_counter = current_counter;
signed_delta = (int32_t)wrapped_delta * encoder_direction_sign;
total_count += signed_delta;
""",
        "16비트 wrap을 이용한 변화량 계산",
    )
    add_table(
        doc,
        ["이전 CNT", "현재 CNT", "실제 이동", "int16 결과"],
        [
            ["100", "112", "정방향 +12", "+12"],
            ["112", "100", "역방향 -12", "-12"],
            ["65530", "5", "wrap 포함 정방향 +11", "+11"],
            ["5", "65530", "wrap 포함 역방향 -11", "-11"],
        ],
        [1900, 1900, 3100, 2460],
        font_size=9.0,
    )
    heading(doc, 2, "카운트를 속도로 바꾸는 공식")
    add_callout(
        doc,
        "공식",
        "속도(m/s) = (이번 주기 count ÷ 1회전 count) × 바퀴 둘레(m) ÷ 주기(s)",
        "note",
    )
    add_body(
        doc,
        "현재 임시 보정은 1600 count/rev, 둘레 0.20106 m, 주기 0.010 s다. "
        "10 ms 동안 11 count라면 (11/1600)×0.20106/0.010 = 약 0.1382 m/s, 즉 138.2 mm/s다.",
    )
    add_table(
        doc,
        ["10 ms delta", "계산 속도", "mm/s", "해석"],
        [
            ["0", "0", "0", "정지 또는 분해능 아래 움직임"],
            ["1", "0.01257 m/s", "12.57", "현재 한 count의 속도 단계"],
            ["10", "0.12566 m/s", "125.66", "텔레메트리에서 자주 보이는 값"],
            ["11", "0.13823 m/s", "138.23", "150 mm/s 목표 근처"],
            ["12", "0.15080 m/s", "150.80", "한 count 차이로 약 12.6 mm/s 점프"],
        ],
        [1700, 2400, 1600, 3660],
        font_size=8.9,
    )
    heading(doc, 2, "보정값은 왜 임시인가")
    add_body(
        doc,
        "1600 count/rev는 벤치에서 손으로 얻은 대략값이고, 64 mm 지름에서 계산한 둘레는 무부하 명목값이다. "
        "타이어 변형, 미끄럼, 기어 유격, 실제 장착 하중 때문에 실차 이동거리와 달라질 수 있다. "
        "여러 바퀴를 평균 내고 실제 바닥 이동거리로 다시 보정해야 한다.",
    )
    add_callout(
        doc,
        "현재 노이즈 설정",
        "TIM4 IC1Filter와 IC2Filter가 모두 0이고 GPIO는 NOPULL이다. 짧고 깨끗한 bench 배선에서는 동작했지만 "
        "긴 케이블·모터 노이즈 환경에서는 필터, 풀업, 차동 수신, 실드 등을 검토해야 한다.",
        "warning",
    )

    # Chapter 7.
    chapter(doc, "7", "pid.c와 feed-forward", "‘틀린 만큼 고친다’를 코드로 표현하기")
    add_figure(
        doc,
        assets["closed_loop"],
        "그림 7-1. 현재 속도를 다시 입력으로 돌려보내는 폐루프",
        "목표 속도와 실제 속도의 오차를 P 제어로 PWM에 보정하는 폐루프 블록",
    )
    heading(doc, 2, "PID의 세 글자")
    add_table(
        doc,
        ["항", "뜻", "직관", "현재 값"],
        [
            ["P", "Proportional", "지금 오차가 크면 크게 고침", "Kp=50"],
            ["I", "Integral", "오래 남은 오차를 누적해 고침", "Ki=0"],
            ["D", "Derivative", "오차 변화가 너무 빠르면 제동", "Kd=0"],
        ],
        [1100, 1800, 4400, 2060],
        font_size=9.0,
    )
    add_body(
        doc,
        "현재는 이름은 PIDController이지만 실제 동작은 P 제어다. Ki와 Kd가 0이기 때문이다. "
        "그래도 적분과 미분, anti-windup 구조를 미리 갖추어 이후 튜닝에 사용할 수 있다.",
    )
    heading(doc, 2, "P 계산의 단위")
    add_body(
        doc,
        "target과 measurement 단위는 m/s이고 출력은 PWM percentage point다. "
        "예를 들어 오차가 0.012 m/s이면 P 보정은 50×0.012=0.6%p다.",
    )
    add_code_block(
        doc,
        "App/Src/pid.c · PID_Update()",
        """
error = target - measurement;
derivative = (error - previous_error) / dt_seconds;
integral_candidate = integral + error * dt_seconds;

unsaturated_output =
    kp * error + ki * integral_candidate + kd * derivative;
output = ClampOutput(pid, unsaturated_output);
""",
        "PID의 교과서 공식을 그대로 읽을 수 있는 부분",
    )
    heading(doc, 2, "왜 22%를 먼저 더하는가")
    add_body(
        doc,
        "벤치 시험에서 약 20% 이하는 정지 마찰을 이기지 못했다. 그래서 비영점 목표가 들어오면 "
        "처음부터 정방향 +22% 또는 역방향 -22%를 주고, PID는 그 주변을 미세 조정한다. "
        "이것이 `ApplySpeedFeedforward()`다.",
    )
    add_table(
        doc,
        ["목표", "측정", "오차", "P 보정", "최종 PWM"],
        [
            ["+0.150", "+0.138", "+0.012", "+0.6", "+22.6%"],
            ["+0.150", "0", "+0.150", "+7.5", "+29.5%"],
            ["-0.150", "-0.138", "-0.012", "-0.6", "-22.6%"],
            ["-0.150", "0", "-0.150", "-7.5", "-29.5%"],
        ],
        [1600, 1600, 1500, 1900, 2760],
        font_size=9.0,
    )
    heading(doc, 2, "왜 과속했다고 갑자기 역회전하지 않는가")
    add_body(
        doc,
        "정방향 목표에서 보정이 너무 음수가 되어도 출력 하한은 0%다. 역방향 목표에서는 상한이 0%다. "
        "즉 현재 목표 방향 안에서만 0~40%를 사용한다. 이것은 과속을 줄이기 위해 갑자기 반대 토크를 거는 것을 막는 보수적 설계다.",
    )
    heading(doc, 2, "Anti-windup은 무엇인가")
    add_body(
        doc,
        "출력이 상한에 붙었는데 적분값이 계속 커지면, 나중에 오차가 줄어도 한동안 과한 출력이 남는다. "
        "현재 코드는 포화 중 적분이 더 나쁜 방향으로 쌓이지 않게 조건부 적분을 한다. Ki=0인 지금은 결과에 영향이 없지만, "
        "나중에 Ki를 켤 때 중요하다.",
    )
    add_callout(
        doc,
        "튜닝 경고",
        "Ki를 바로 크게 올리면 저속 양자화, 마찰, 40% 제한과 상호작용해 진동·오버슈트가 생길 수 있다. "
        "고정된 시험 조건에서 작은 값부터 단계 응답으로 검증해야 한다.",
        "warning",
    )

    # Chapter 8.
    chapter(doc, "8", "10 ms 제어와 안전 게이트", "task_control.c · task_safety.c · safety.c")
    add_figure(
        doc,
        assets["rtos"],
        "그림 8-1. Safety, Control, Comm, Ultrasonic, Telemetry의 서로 다른 시간축",
        "FreeRTOS 안전 제어 통신 초음파 태스크와 텔레메트리의 주기를 비교한 시간선",
    )
    heading(doc, 2, "Control_Init()")
    add_body(
        doc,
        "공유 명령과 공유 상태를 0으로 지우고 모드를 DISABLED로 설정한다. 안전 요청은 처음부터 "
        "`STOP + SENSOR_INIT + block_reverse=true`로 둔다. 센서가 정상임을 증명하기 전에는 모터를 허가하지 않는 "
        "fail-safe 기본값이다. safety request 전용 Mutex도 만든다. 그 뒤 fault를 비우고 "
        "PID를 Kp=50, Ki=0, Kd=0, 출력 보정 범위 -18~+18로 초기화한다. "
        "18은 최대 40%에서 기본 22%를 뺀 값이다.",
    )
    heading(doc, 2, "Control_Task() 시작 시 한 번만 하는 일")
    add_numbered_list(doc, [
        "Motor_Init()으로 두 PWM 채널을 시작한다.",
        "Encoder_Init()으로 TIM4 encoder counter를 시작한다.",
        "하나라도 실패하면 HARDWARE_INIT fault를 latch하고 Motor_Disable()한다.",
        "다음 wake tick을 현재 RTOS tick으로 저장한다.",
    ])
    heading(doc, 2, "매 10 ms 루프의 정확한 순서")
    add_table(
        doc,
        ["순서", "코드", "초보자 번역"],
        [
            ["1", "latched fault 읽기", "이전에 치명적인 실패가 있었나?"],
            ["2", "Encoder_Update()", "지금 속도를 잰다"],
            ["3", "CopyCommand()", "가장 최신 목표를 복사한다"],
            ["4", "CopySafetyRequest()", "SafetyTask의 최신 허가·정지 판단을 복사한다"],
            ["5", "50 ms freshness 검사", "안전 판단이 오래됐으면 고장으로 보고 정지한다"],
            ["6", "mode change 검사", "모드가 바뀌면 PID와 모터를 리셋한다"],
            ["7", "hardware/safety gate", "고장·정지 요청·래치·stale이면 출력하지 않는다"],
            ["8", "switch(mode)", "허가됐을 때만 Open loop / Speed PID / Disabled 실행"],
            ["9", "ControlState 저장", "목표·속도·PWM·fault를 통신 태스크에 공개한다"],
            ["10", "osDelayUntil()", "정확한 10 ms 박자로 다음 루프를 기다린다"],
        ],
        [850, 2820, 5690],
        font_size=8.7,
    )
    add_figure(
        doc,
        assets["safety"],
        "그림 8-2. 초음파 한 번의 측정이 안전 상태와 모터 차단으로 이어지는 길",
        "초음파 거리의 중앙값 필터와 안전 임계값이 ControlTask의 출력 허가 또는 차단으로 이어지는 흐름",
    )
    heading(doc, 2, "안전 게이트가 제어보다 먼저 결정한다")
    add_body(
        doc,
        "SafetyTask는 10 ms마다 초음파 상태, 현재 명령 방향, 측정된 이동 방향, 실제 모터 출력 0 여부를 모아 "
        "`Safety_Evaluate()`에 넣는다. 결과를 `Control_SetSafetyRequest()`로 전달하면 ControlTask는 명령의 크기보다 "
        "안전 요청을 먼저 검사한다. `stop_request`이면 일반 Disable, `latched`이면 EmergencyDisable을 호출한다.",
    )
    add_code_block(
        doc,
        "App/Src/task_control.c · safety gate",
        """
else if (safety_request.stop_request ||
         safety_request.latched ||
         safety_request_stale)
{
    PID_Reset(&speed_pid);
    if (safety_request.latched) Motor_EmergencyDisable();
    else                        Motor_Disable();
}
else
{
    switch (command.mode) { /* 허가된 경우에만 제어 */ }
}
""",
        "통신 명령이 있어도 안전 허가가 없으면 PWM 경로에 들어가지 않는다",
    )
    heading(doc, 2, "거리 임계값과 히스테리시스")
    add_table(
        doc,
        ["조건", "상태", "출력", "해제 조건"],
        [
            ["거리 ≥ 0.65 m", "CLEAR", "전진 허가", "계속 정상 측정"],
            ["거리 < 0.60 m", "CAUTION", "현재 코드상 전진 허가", "0.65 m 이상"],
            ["거리 < 0.20 m", "STOP", "일반 Disable", "0.30 m 이상 정상 샘플 3회"],
            ["거리 < 0.20 m", "E-STOP 래치", "EmergencyDisable", "정지·정상 센서·0.60 m 이상에서 reset"],
        ],
        [2500, 2050, 2400, 2410],
        font_size=8.5,
    )
    add_body(
        doc,
        "진입값과 해제값을 다르게 둔 것을 히스테리시스라고 한다. 0.20 m 근처의 측정 잡음 때문에 STOP과 CLEAR가 "
        "계속 흔들리는 현상을 막는다. 안전 상태 머신은 센서 데이터가 200 ms보다 오래되거나, 전진 중 timeout·범위 오류가 "
        "3회 누적되면 비상 래치까지 올릴 수 있다. ControlTask는 별도로 SafetyTask 요청이 50 ms보다 오래되면 즉시 정지한다.",
    )
    add_callout(
        doc,
        "후진은 현재 의도적으로 차단",
        "후방 거리 센서가 없기 때문에 `block_reverse=true`가 기본이고, 음수 PWM·음수 속도 명령 또는 실제 후진 움직임은 "
        "`REVERSE_UNPROTECTED` STOP을 만든다. 따라서 현재 펌웨어에서 후진 진단이 움직이지 않는 것은 정상 안전 동작이다.",
        "warning",
    )
    heading(doc, 2, "모드가 바뀔 때 반드시 정지하는 이유")
    add_body(
        doc,
        "Open loop에서 Speed PID로 바꾸는 순간, 이전 모드의 PWM과 PID 기억을 그대로 이어 쓰면 예상치 못한 출력이 나올 수 있다. "
        "현재 코드는 mode가 달라지면 `PID_Reset()`과 `Motor_Disable()`을 먼저 호출한 뒤 새 모드를 수행한다.",
    )
    heading(doc, 2, "세 모드의 동작")
    add_code_block(
        doc,
        "App/Src/task_control.c · Control_Task() switch 요약",
        """
OPEN_LOOP:
    Motor_Enable();
    Motor_SetPercent(command.pwm_percent);

SPEED_PID:
    Encoder 확인 → PID_Update → feed-forward → Motor_SetPercent();

DISABLED:
    PID_Reset();
    Motor_Disable();
""",
        "실제 코드를 초보자용으로 압축한 구조",
    )
    heading(doc, 2, "shared_command와 shared_state는 왜 critical section으로 감싸나")
    add_body(
        doc,
        "CommTask가 명령을 쓰는 순간 ControlTask가 같은 구조체를 읽으면, 일부 필드만 새 값인 찢어진 데이터가 될 수 있다. "
        "`taskENTER_CRITICAL()`은 아주 짧은 복사 동안 선점을 막아 구조체를 한 덩어리로 주고받게 한다.",
    )
    add_callout(
        doc,
        "현재 구조의 특징",
        "Queue나 Mutex를 별도로 만들지 않고 짧은 critical section으로 최신값 하나만 공유한다. "
        "주행 명령은 오래된 명령을 모두 처리하기보다 가장 최신 목표가 중요하므로 단순하고 합리적인 선택이다.",
        "note",
    )
    heading(doc, 2, "fault가 active와 latched로 나뉘는 이유")
    add_body(
        doc,
        "하드웨어 초기화 실패는 재부팅 없이 단순 확인만으로 복구할 수 없으므로 latched 상태를 유지한다. "
        "엔코더 보정 부족 같은 조건은 현재 상태에 따라 active fault로 켜졌다 꺼질 수 있다. "
        "`Control_ClearFaults()`도 HARDWARE_INIT 비트는 지우지 못하게 막는다.",
    )

    # Chapter 9.
    chapter(doc, "9", "app_freertos.c와 UART 진단 경로", "PC 명령을 안전하게 ControlTask로 전달하기")
    heading(doc, 2, "실제로 만들어지는 태스크")
    add_table(
        doc,
        ["태스크", "우선순위", "주기", "주요 일"],
        [
            ["SafetyTask", "High", "10 ms", "거리·방향·sensor stale 판단, 안전 요청, E-STOP"],
            ["ControlTask", "AboveNormal", "10 ms", "엔코더, PID, 모터 출력, 상태 저장"],
            ["CommTask", "Normal", "1 ms", "UART parser, 명령 적용, telemetry"],
            ["UltrasonicTask", "Normal", "60 ms", "TRIG, ECHO 대기, 거리 계산, 중앙값 필터"],
            ["defaultTask", "Normal", "1 ms delay", "CubeMX 기본 빈 태스크"],
        ],
        [1900, 1800, 1500, 4160],
        font_size=8.6,
    )
    add_body(
        doc,
        "SafetyTask가 가장 높고 ControlTask가 그 다음이라, 안전 판단과 출력 차단이 통신 처리보다 먼저 실행될 수 있다. "
        "UART 수신은 DMA와 interrupt가 바이트 이동을 담당하고, 파싱과 명령 처리는 CommTask 문맥에서 한다.",
    )
    heading(doc, 2, "초음파 한 번을 재는 과정")
    add_numbered_list(doc, [
        "UltrasonicTask가 PA5 TRIG를 10 µs 동안 HIGH로 만든다.",
        "TIM2 CH2가 PB3 ECHO의 rising edge 시간을 저장한다.",
        "같은 채널의 극성을 falling으로 바꾸고 다음 edge 시간을 저장한다.",
        "32비트 뺄셈으로 펄스 폭을 구한다. TIM2는 1 MHz라 1 count=1 µs다.",
        "30 ms 안에 결과가 없으면 TIMEOUT, 0.03~4.00 m 밖이면 OUT_OF_RANGE로 기록한다.",
        "정상 거리 최근 3개의 중앙값을 저장하고, 60 ms 뒤 다음 측정을 시작한다.",
    ])
    add_callout(
        doc,
        "거리 공식",
        "거리(m) = ECHO 펄스 폭(µs) × 0.0001715. 소리가 왕복하므로 음속 약 343 m/s를 2로 나눈 값이다.",
        "note",
    )
    heading(doc, 2, "진단 Motor Test 경로")
    add_numbered_list(doc, [
        "PC 도구가 DIAG_MOTOR_TEST_REQUEST(0xF2)를 보낸다.",
        "CommService가 길이, ±400 permille, 100~2000 ms 범위를 검사한다.",
        "응답 프레임을 TX queue에 넣는 데 성공해야만 motor test를 pending으로 만든다.",
        "CommTask가 pending 명령을 꺼내 Control_SetOpenLoopPercent()를 호출한다.",
        "지정 시간이 지나면 펌웨어가 Control_SetDisabled()를 호출한다.",
        "PC 도구도 종료 시 별도 stop을 보내 이중 차단한다.",
    ])
    add_callout(
        doc,
        "왜 응답을 먼저 보내나",
        "PC가 ‘승인됨’ 응답을 받지 못한 채 모터만 움직이는 상황을 막기 위해서다. "
        "통신 응답 queue가 가득 차면 모터 상태를 바꾸지 않는다.",
        "tip",
    )
    add_callout(
        doc,
        "진단 명령도 안전 규칙을 우회하지 않는다",
        "응답 ACK와 타이머가 정상이어도 초음파가 INIT·TIMEOUT·STOP이면 모터는 움직이지 않는다. "
        "후진 진단은 후방 센서가 없는 현재 정책에서 REVERSE_UNPROTECTED로 차단된다.",
        "warning",
    )
    heading(doc, 2, "PID Test 경로")
    add_body(
        doc,
        "DIAG_PID_TEST_REQUEST(0xF4)는 ±300 mm/s 이내, 500~5000 ms 범위의 비영점 목표만 받는다. "
        "CommTask가 `Control_SetSpeedTarget()`으로 바꾸고 제한 시간이 지나면 자동 정지한다. "
        "이 경로는 생산용 drive limit을 풀지 않고도 폐루프를 시험할 수 있게 만든 별도 진단 통로다.",
    )
    heading(doc, 2, "생산용 CMD_DRIVE가 지금 움직이지 않는 이유")
    add_code_block(
        doc,
        "App/Src/app_freertos.c",
        """
#define COMM_MAX_ABS_SPEED_MM_S     0
#define COMM_MAX_ABS_STEERING_CDEG  0
""",
        "실측 한계를 넣기 전 비영점 생산 명령을 잠그는 값",
    )
    add_body(
        doc,
        "`CommService_SetDriveLimits(0, 0)`은 limits configured 상태를 false로 만든다. "
        "따라서 비영점 CMD_DRIVE는 거부된다. 이것은 미완성 버그가 아니라 실차 한계를 재기 전 의도적 안전 잠금이다.",
    )
    heading(doc, 2, "통신 timeout")
    add_body(
        doc,
        "유효한 drive 명령이 300 ms 넘게 오지 않으면 neutral command를 만들고 재출발 허용 상태를 잠근다. "
        "이후에는 먼저 drive_enable=false, 속도=0, 조향=0인 중립 명령을 받아야 재무장된다. "
        "단, 시간 제한 진단은 자체 stop tick으로 별도 관리된다.",
    )
    heading(doc, 2, "Telemetry가 보여주는 것")
    add_table(
        doc,
        ["필드", "ControlState 원본", "단위 변환", "초보자가 볼 것"],
        [
            ["target_speed_mm_s", "target_speed_mps", "×1000", "명령이 제대로 들어왔나"],
            ["measured_speed_mm_s", "measured_speed_mps", "×1000", "엔코더와 부호가 맞나"],
            ["encoder_count", "total_count", "int64→int32 포화", "정방향 +, 역방향 -인가"],
            ["motor_duty_permille", "applied_pwm_percent", "×10", "PWM이 부하에 반응하나"],
            ["active_fault_bits", "fault flags", "비트 매핑", "왜 READY/FAULT인지"],
        ],
        [2200, 2500, 1850, 2810],
        font_size=8.7,
    )
    add_body(
        doc,
        "초음파 INIT·TIMEOUT·OUT_OF_RANGE·STALE은 통신의 `RANGE_LOST` fault로, 안전 래치는 "
        "`ESTOP_ACTIVE`로 변환된다. 래치는 아니지만 `stop_request=true`이면 telemetry state는 `SAFE_STOP`이다. "
        "따라서 ‘명령은 들어왔는데 duty가 0’이면 target만 보지 말고 state와 fault를 함께 읽어야 한다.",
    )

    part_page(
        doc,
        "3",
        "손으로 계산하고 안전하게 시험한다",
        "이제 코드를 바라보는 데서 끝내지 않는다. 한 번의 제어 주기를 숫자로 추적하고, 실제 시험에서 성공과 실패를 구분한다.",
    )

    # Chapter 10.
    chapter(doc, "10", "150 mm/s 명령 한 바퀴 추적", "실제 숫자로 ControlTask 한 주기를 따라가기")
    heading(doc, 2, "상황 설정")
    for item in [
        "PID 진단 목표: +150 mm/s = +0.150 m/s",
        "안전 전제: 초음파 상태 정상, 거리 0.65 m 이상, 최신 SafetyRequest",
        "이번 10 ms 동안 엔코더 변화량: +11 count",
        "보정값: 1600 count/rev, wheel circumference=0.20106 m",
        "Kp=50, Ki=0, Kd=0",
        "기본 PWM=22%, 최대 PWM=40%",
    ]:
        add_bullet(doc, item)
    heading(doc, 2, "0단계: 안전 허가")
    add_body(
        doc,
        "ControlTask는 아래 계산을 무조건 실행하지 않는다. SafetyTask가 `stop_request=false`, `latched=false`인 최신 요청을 "
        "보내야만 PID와 PWM switch로 들어간다. 센서가 아직 INIT이거나 timeout이면 계산상 22.6%가 나와도 실제 출력은 0%다.",
    )
    heading(doc, 2, "1단계: 엔코더 속도")
    add_body(
        doc,
        "회전수 = 11/1600 = 0.006875 rev. 이동거리 = 0.006875×0.20106 = 0.0013823 m. "
        "이를 0.010 s로 나누면 0.13823 m/s다.",
    )
    heading(doc, 2, "2단계: 오차")
    add_body(doc, "error = target - measurement = 0.15000 - 0.13823 = 0.01177 m/s.")
    heading(doc, 2, "3단계: P 보정")
    add_body(doc, "pid_correction = Kp × error = 50 × 0.01177 = 0.5885 percentage point.")
    heading(doc, 2, "4단계: feed-forward")
    add_body(doc, "최종 명령 = 22 + 0.5885 = 약 22.59%. 코드의 float 연산과 telemetry 반올림에 따라 22.6%로 보인다.")
    heading(doc, 2, "5단계: CCR")
    add_body(doc, "CCR = round(500 × 22.5885 / 100) = round(112.9425) = 113. TIM5 CH1의 CCR1에 113을 쓴다.")
    add_table(
        doc,
        ["단계", "값", "다음으로 전달"],
        [
            ["Encoder_Update", "+11 count → 0.13823 m/s", "encoder_sample.speed_mps"],
            ["PID_Update", "error 0.01177 → +0.5885", "pid_correction"],
            ["ApplySpeedFeedforward", "22 + 0.5885 → 22.5885%", "applied_pwm"],
            ["Motor_SetPercent", "22.5885% → CCR1 113", "TIM5_CH1"],
            ["SendTelemetryIfDue", "22.6% → 226 permille", "UART TELEMETRY_DRIVE"],
        ],
        [2200, 3500, 3660],
        font_size=8.9,
    )
    heading(doc, 2, "바퀴가 완전히 멈추면")
    add_body(
        doc,
        "measurement=0이면 error=0.150, P 보정=7.5, 최종 PWM=29.5%다. "
        "따라서 현재 P-only 설계는 목표 150 mm/s에서 stall이 나도 40%까지 쓰지 않는다. "
        "지속 부하에서 목표 복구가 부족했던 실험 결과와 일치한다.",
    )
    add_callout(
        doc,
        "핵심 통찰",
        "코드가 ‘최대 40%’라고 해도 특정 목표와 Kp 조합에서는 그 상한에 도달하지 않을 수 있다. "
        "한계값과 실제 계산 결과를 구분해야 한다.",
        "tip",
    )
    question_box(
        doc,
        "같은 조건에서 엔코더 delta가 12 count라면 P 보정은 양수일까 음수일까?",
        "12 count는 약 150.8 mm/s로 목표보다 조금 빠르다.",
    )

    # Chapter 11.
    chapter(doc, "11", "안전하게 실행해 보는 실습", "명령, telemetry, 자동 정지 확인")
    heading(doc, 2, "실습 0: 빌드와 연결만 확인")
    add_code_block(
        doc,
        "PowerShell · 프로젝트 루트",
        r"""
py tools\uart_protocol_test.py self-test
py tools\uart_protocol_test.py list
py tools\uart_protocol_test.py echo --port COM14 --text STM32
""",
        "모터 전원 없이도 가능한 통신 확인",
    )
    add_body(
        doc,
        "self-test는 PC 도구 내부의 frame/CRC/parser를 확인한다. list로 실제 COM 번호를 찾고, "
        "echo 응답이 오면 USART1, ST-LINK VCP, DMA, parser, TX queue가 기본적으로 연결된 것이다.",
    )
    add_callout(
        doc,
        "모터 시험 전 초음파가 먼저",
        "PA5 TRIG, PB3 ECHO, 공통 GND를 확인하고 ECHO가 STM32 입력에서 3.3 V를 넘지 않게 한다. "
        "센서가 INIT·TIMEOUT이면 fail-safe로 모든 전진 실습도 PWM 0이 정상이다.",
        "danger",
    )
    heading(doc, 2, "실습 1: 엔코더만 손으로 돌리기")
    add_code_block(
        doc,
        "PowerShell",
        r"""
py tools\uart_protocol_test.py encoder-monitor --port COM14 --seconds 10
""",
        "모터 전원을 끈 상태에서 바퀴를 손으로 한 바퀴",
    )
    add_table(
        doc,
        ["관찰", "정상", "비정상일 때"],
        [
            ["정방향 count", "증가", "A/B 교환 또는 direction sign 검토"],
            ["역방향 count", "감소", "부호 설정 검토"],
            ["정지 중 count", "거의 고정", "노이즈·접지·필터 점검"],
            ["한 바퀴 절댓값", "반복값이 비슷함", "기계 유격·신호 손실 점검"],
        ],
        [2100, 2700, 4560],
        font_size=9.0,
    )
    heading(doc, 2, "실습 2: 짧은 open-loop")
    add_code_block(
        doc,
        "PowerShell",
        r"""
py tools\uart_protocol_test.py encoder-test --port COM14 --percent 25 --duration 500
""",
        "바퀴를 띄우고 전방 장애물을 0.65 m보다 멀리 둔 상태에서 0.5초 정방향 시험",
    )
    add_callout(
        doc,
        "정상 기준",
        "정방향 count 증가, 명령 중 양수 duty, 명령 시간이 끝난 뒤 PWM=0과 READY 상태로 돌아와야 한다.",
        "tip",
    )
    add_callout(
        doc,
        "음수 진단은 회전 시험이 아니다",
        "현재 후방 센서가 없어 음수 open-loop/PID 명령은 `REVERSE_UNPROTECTED`로 차단된다. "
        "음수 명령에서 duty=0, SAFE_STOP, 바퀴 정지가 보이면 안전 정책이 정상이다. 도구가 ‘움직임 없음’을 FAIL로 표시해도 "
        "이는 모터 고장이 아니라 시험 도구의 예전 기대값과 현재 안전 정책의 차이다.",
        "warning",
    )
    heading(doc, 2, "실습 3: PID 진단")
    add_code_block(
        doc,
        "PowerShell",
        r"""
py tools\uart_protocol_test.py pid-test --port COM14 --target-mm-s 150 --duration 3000
""",
        "거리 CLEAR 상태에서 현재 임시 보정값을 사용하는 정방향 무부하 bench 시험",
    )
    heading(doc, 2, "기록할 항목")
    for item in [
        "평균 목표 속도와 평균 측정 속도",
        "최대 PWM과 마지막 PWM",
        "엔코더 총 변화량과 부호",
        "자동 정지 여부",
        "비정상 소음·진동·모듈 온도",
        "배터리 전압과 바퀴/차체 상태",
    ]:
        add_bullet(doc, item)
    add_callout(
        doc,
        "실험 원칙",
        "한 번에 한 변수만 바꾼다. Kp, 목표 속도, 하중을 동시에 바꾸면 무엇이 결과를 바꿨는지 알 수 없다.",
        "note",
    )
    heading(doc, 2, "실습 4: 안전 정지 확인")
    add_numbered_list(doc, [
        "바퀴를 띄운 채 전방 거리를 0.65 m 이상 확보하고 +25% 진단을 시작한다.",
        "반사판을 천천히 0.20 m 안쪽으로 이동한다. 손·신체를 바퀴나 구동계에 넣지 않는다.",
        "PWM=0, Enable=LOW, telemetry SAFE_STOP이 되는지 확인한다.",
        "반사판을 0.30 m 이상으로 옮기고 정상 측정 3회 뒤에만 재허가되는지 확인한다.",
        "센서 ECHO 선을 분리하는 시험은 별도 저전력 bench에서 수행하고 RANGE_LOST와 정지를 확인한다.",
    ])

    # Chapter 12.
    chapter(doc, "12", "모터가 안 돌 때의 디버깅 순서", "증상에서 파일과 핀으로 거슬러 올라가기")
    heading(doc, 2, "가장 좋은 순서: 전원 → 핀 → 타이머 → 드라이버 → 제어")
    add_table(
        doc,
        ["증상", "첫 확인", "그다음", "관련 코드"],
        [
            ["모터가 항상 안 돎", "SAFE_STOP/RANGE_LOST", "초음파 상태·SafetyRequest 최신성", "task_safety, task_ultrasonic"],
            ["안전 CLEAR인데 안 돎", "모터 전원·GND·Enable", "TIM5 PWM start/fault", "Motor_Init, main GPIO"],
            ["정방향만 됨", "PA3 실제 파형·LPWM 배선", "CH4 start/CCR4", "motor.c, MSP"],
            ["후진 명령이 안 움직임", "현재 정상: reverse 차단", "REVERSE_UNPROTECTED", "safety.c, task_safety.c"],
            ["모터는 도는데 속도 0", "PB6/PB7 파형", "TIM4 CNT 변화", "encoder.c, TIM4 init"],
            ["속도 부호 반대", "A/B 순서", "direction sign", "Encoder_SetDirectionSign"],
            ["PWM은 증가하나 속도 회복 실패", "전원·하중·마찰", "Kp/Ki와 상한", "pid.c, task_control.c"],
            ["CMD_DRIVE 거부", "drive limit=0", "진단 명령 사용", "app_freertos.c, comm_service.c"],
            ["몇 초 뒤 멈춤", "진단 duration/timeout", "fault bits", "ApplyPendingCommands"],
        ],
        [2200, 2450, 2400, 2310],
        font_size=8.2,
    )
    heading(doc, 2, "1단계: 소프트웨어 상태만 보기")
    for item in [
        "ControlState.mode가 무엇인가?",
        "fault_flags에 HARDWARE_INIT 또는 ENCODER_NOT_CALIBRATED가 있는가?",
        "telemetry state가 SAFE_STOP/FAULT인가? RANGE_LOST·ESTOP_ACTIVE가 있는가?",
        "UltrasonicState가 OK이며 SafetyRequest가 50 ms 이내 최신값인가?",
        "applied_pwm_percent가 0인가, 0이 아닌가?",
        "encoder_total_count가 손으로 돌릴 때 변하는가?",
        "CommService 통계에서 invalid payload, rejection, timeout이 증가하는가?",
    ]:
        add_bullet(doc, item)
    heading(doc, 2, "2단계: 오실로스코프/로직 애널라이저")
    add_body(
        doc,
        "PA0과 PA3에서 같은 20 kHz 주파수가 보이는지 확인한다. 양수 명령일 때 PA0만 Duty가 있고 PA3는 0, "
        "음수 명령일 때는 반대여야 한다. PE2/PE3는 Enable 중 HIGH, Disabled 중 LOW여야 한다.",
    )
    add_callout(
        doc,
        "주의",
        "오실로스코프 GND 클립을 모터 전원 고측이나 잘못된 노드에 연결하면 단락이 날 수 있다. "
        "측정 장비의 접지 구조를 모르면 전원을 넣기 전에 지도자에게 확인한다.",
        "danger",
    )
    heading(doc, 2, "초음파가 계속 TIMEOUT이면")
    for item in [
        "PA5에서 약 10 µs HIGH TRIG가 60 ms마다 보이는가?",
        "PB3 ECHO가 TIM2_CH2 AF1로 설정됐고 TIM2 IRQ가 priority 6으로 활성화됐는가?",
        "센서 ECHO 전압이 3.3 V 이하인가? 5 V형 센서면 분압기·레벨 시프터를 먼저 넣는다.",
        "보드 SB9가 연결돼 PB3의 SWO와 전기적으로 다투지 않는지 실물을 확인한다.",
        "ECHO rising/falling이 30 ms 안에 오며 0.03~4.00 m 범위인가?",
    ]:
        add_bullet(doc, item)
    heading(doc, 2, "3단계: 엔코더 노이즈 찾기")
    for item in [
        "모터 전원 OFF인데 count가 변하면 신호선/풀업/접지 문제 가능성이 크다.",
        "모터가 돌 때만 튄다면 모터 EMI, 공통 GND bounce, 케이블 배치, 입력 필터를 의심한다.",
        "A/B 중 한 채널만 보이면 방향 판별과 4배 계수가 무너진다.",
        "10 ms delta 절댓값이 32767을 넘을 가능성이 있으면 현재 wrap 계산 전제가 깨진다.",
    ]:
        add_bullet(doc, item)
    heading(doc, 2, "Error_Handler에 들어갔을 때")
    add_body(
        doc,
        "클록 설정, GPIO, TIM, UART, DMA 초기화 중 HAL_OK가 아닌 결과가 나오면 Error_Handler로 간다. "
        "디버거 breakpoint를 Error_Handler와 Motor_Init 실패 분기에 걸고 call stack과 HAL handle의 ErrorCode를 본다. "
        "빌드 성공은 이 런타임 초기화 성공을 보장하지 않는다.",
    )

    # Chapter 13.
    chapter(doc, "13", "다음 코드를 안전하게 바꾸는 법", "작은 변경, 충돌 검사, 검증 기록")
    heading(doc, 2, "변경 전 반드시 답할 질문")
    for item in [
        "바꾸려는 핀은 `.ioc`와 공식 보드 회로도에서 정말 비어 있는가?",
        "그 핀이 원하는 Timer Channel과 Alternate Function을 실제로 지원하는가?",
        "같은 Timer의 다른 채널이 기존 기능과 주파수를 공유해도 되는가?",
        "APB 타이머 입력 클록을 기준으로 PSC/ARR을 다시 계산했는가?",
        "CubeMX 재생성 뒤에도 사용자 코드가 보존되는가?",
        "실패 시 모터를 0으로 만들고 Enable을 내리는 경로가 있는가?",
    ]:
        add_bullet(doc, item)
    heading(doc, 2, "예제 1: Kp만 바꾸기")
    add_body(
        doc,
        "`CONTROL_PID_KP`를 한 단계만 바꾸고 같은 배터리, 같은 바퀴, 같은 목표 속도에서 반복 측정한다. "
        "상승시간, overshoot, 정상상태 오차, 최대 PWM을 표로 남긴다. 결과가 나쁘면 즉시 원래 값으로 돌아갈 수 있게 변경 전 값을 기록한다.",
    )
    add_table(
        doc,
        ["시험", "Kp", "목표", "평균 속도", "최대 PWM", "진동/오버슈트"],
        [
            ["기준", "50", "150 mm/s", "기록", "기록", "기록"],
            ["변경 A", "예: 60", "150 mm/s", "기록", "기록", "기록"],
            ["변경 B", "예: 40", "150 mm/s", "기록", "기록", "기록"],
        ],
        [1100, 1100, 1750, 1900, 1600, 1910],
        font_size=8.7,
    )
    heading(doc, 2, "예제 2: 작은 Ki 추가")
    add_body(
        doc,
        "지속 하중의 정상상태 오차를 줄이기 위해 작은 Ki를 시험할 수 있다. 하지만 먼저 encoder 속도 노이즈와 "
        "출력 포화 동작을 기록해야 한다. Ki를 켜면 anti-windup이 실제로 의미를 가지며, stop·mode change에서 "
        "`PID_Reset()`이 적분값을 지우는지 확인해야 한다.",
    )
    heading(doc, 2, "예제 3: 엔코더 필터 추가")
    add_body(
        doc,
        "CubeMX에서 TIM4 IC1Filter/IC2Filter를 같은 값으로 설정하고 `.ioc`와 `main.c` 생성 결과를 함께 확인한다. "
        "필터는 노이즈를 줄이지만 너무 크면 빠른 edge를 놓칠 수 있다. 정지 노이즈와 최대 속도 count 손실을 모두 시험해야 한다.",
    )
    heading(doc, 2, "현재 프로젝트의 우선 개선 순서")
    add_table(
        doc,
        ["우선", "개선", "이유", "완료 기준"],
        [
            ["1", "PA0/PA3/PB6 공유핀 재검토", "보드 내장 기능과 물리 충돌", "회로도 근거로 독립 핀 확정"],
            ["2", "실차 encoder 보정", "속도 계산의 기준", "여러 회전·거리 평균값 확보"],
            ["3", "Stall/encoder 단선 fault", "현재 fault 범위가 제한적", "고장 주입 시 자동 정지"],
            ["4", "반복 가능한 부하 시험", "손으로 누르는 시험은 재현성 낮음", "동일 조건 단계 응답 비교"],
            ["5", "Kp/Ki 튜닝", "지속 부하 오차", "상승·정착·오차 기준 만족"],
            ["6", "생산 drive limit 설정", "현재 의도적으로 0", "실측 안전 한계 승인"],
        ],
        [1000, 2600, 2900, 2860],
        font_size=8.5,
    )
    add_callout(
        doc,
        "좋은 임베디드 변경",
        "코드 한 줄이 아니라 `.ioc → 생성 초기화 → BSP → Task → 시험 절차 → 정상/실패 기준`이 한 묶음으로 움직이는 변경이다.",
        "tip",
    )

    # Appendices.
    heading(doc, 1, "부록 A · 공식 치트시트", page_break_before=True)
    add_table(
        doc,
        ["주제", "공식", "현재 값 예"],
        [
            ["PWM 주파수", "TimerCLK/(PSC+1)/(ARR+1)", "90 MHz/9/500=20 kHz"],
            ["Duty→CCR", "(ARR+1)×Duty/100", "500×30%=150"],
            ["초음파 거리", "pulse_us×0.0001715", "2000 µs≈0.343 m"],
            ["엔코더 회전수", "delta/counts_per_rev", "11/1600=0.006875 rev"],
            ["이동거리", "rev×circumference", "0.006875×0.20106=0.001382 m"],
            ["속도", "distance/dt", "0.001382/0.010=0.1382 m/s"],
            ["P 보정", "Kp×(target-measurement)", "50×0.01177=0.5885%p"],
            ["최종 정방향 PWM", "22%+P correction", "약 22.59%"],
        ],
        [1800, 3700, 3860],
        font_size=9.0,
    )

    heading(doc, 2, "단위 변환")
    for item in [
        "1 m/s = 1000 mm/s",
        "150 mm/s = 0.150 m/s",
        "1% = 10 permille",
        "22.6% = 226 permille",
        "10 ms = 0.010 s",
        "20 kHz의 한 주기 = 50 µs",
    ]:
        add_bullet(doc, item)

    heading(doc, 1, "부록 B · 초보자 용어사전", page_break_before=True)
    glossary = [
        ("AF (Alternate Function)", "GPIO 핀을 단순 입출력이 아니라 TIM/UART 같은 주변장치에 연결하는 모드."),
        ("ARR (Auto-Reload Register)", "타이머가 어디까지 센 뒤 0으로 돌아갈지 정하는 값."),
        ("BSP", "Board Support Package. 특정 보드/배선의 하드웨어를 감싸는 드라이버 층."),
        ("CCR (Capture/Compare Register)", "PWM에서 HIGH가 유지될 count 경계를 정하는 값."),
        ("Critical section", "아주 짧은 동안 다른 태스크/인터럽트가 공유 데이터에 끼어들지 못하게 하는 구간."),
        ("DMA", "CPU가 한 바이트씩 옮기지 않아도 주변장치와 메모리 사이 데이터를 옮기는 하드웨어."),
        ("Duty", "한 PWM 주기 중 HIGH인 시간의 비율."),
        ("Encoder", "회전량과 방향을 펄스로 알려주는 센서."),
        ("Feed-forward", "오차가 생기기 전부터 필요한 것으로 예상되는 기본 출력을 먼저 주는 방법."),
        ("HAL", "ST가 제공하는 하드웨어 추상화 함수 모음."),
        ("Handle", "htim5, huart1처럼 주변장치 설정과 상태를 담는 구조체."),
        ("ISR/IRQ", "하드웨어 사건이 생겼을 때 즉시 실행되는 interrupt 처리 흐름."),
        ("Hysteresis", "상태 진입과 해제 경계를 다르게 두어 임계값 근처에서 상태가 흔들리는 것을 막는 방법."),
        ("Latch", "한 번 생기면 조건이 사라져도 명시적 복구 전까지 기억하는 상태."),
        ("Median filter", "최근 값들을 정렬해 가운데 값을 고르는 필터. 한 번의 큰 튐에 강하다."),
        ("Open loop", "실제 결과를 보지 않고 명령만 출력하는 제어."),
        ("Closed loop", "측정 결과를 되먹임하여 명령을 계속 보정하는 제어."),
        ("PID", "비례·적분·미분 항으로 오차를 줄이는 제어기."),
        ("Prescaler", "빠른 타이머 클록을 일정 비율로 나누는 값."),
        ("Quadrature", "90도 위상차가 있는 A/B 두 신호로 회전 방향까지 알아내는 방식."),
        ("RTOS", "여러 태스크의 우선순위와 실행 시점을 관리하는 운영체제."),
        ("Telemetry", "장치 내부 상태를 외부로 주기적으로 보내는 데이터."),
        ("Wrap-around", "카운터 최댓값 다음이 0으로 돌아가는 현상."),
    ]
    for term, definition in glossary:
        add_body(doc, f"{term} — {definition}", bold_lead=term)

    heading(doc, 1, "부록 C · 연습문제와 해설", page_break_before=True)
    heading(doc, 2, "문제를 풀고 바로 해설 확인하기")
    questions = [
        ("1", "TIM5 입력 90 MHz, PSC=8, ARR=999라면 PWM 주파수는?", "90,000,000/9/1000=10,000 Hz, 즉 10 kHz."),
        ("2", "30% 명령에서 ARR=499일 때 CCR은?", "(499+1)×0.30=150."),
        ("3", "Motor_SetPercent(-30)는 어느 채널을 쓰나?", "TIM5 CH4, 즉 PA3 LPWM에 CCR 150을 쓴다."),
        ("4", "엔코더가 65530에서 5가 되면 delta는?", "정방향 +11. uint16 뺄셈과 int16 해석으로 wrap을 복원한다."),
        ("5", "왜 모드가 바뀔 때 Motor_Disable을 먼저 하나?", "이전 모드의 PWM과 PID 기억이 새 모드로 튀는 것을 막기 위해서다."),
        ("6", "Kp=50에서 오차 0.02 m/s의 P 보정은?", "1.0 percentage point."),
        ("7", "생산용 CMD_DRIVE가 비영점 명령을 거부하는 이유는?", "실측 speed/steering limit이 아직 0이라 의도적으로 잠겨 있기 때문이다."),
        ("8", "PA0 PWM 중 B1 버튼을 누르면 안 되는 이유는?", "보드에서 B1이 PA0를 GND에 연결하므로 출력 핀과 충돌할 수 있다."),
        ("9", "측정 속도가 0인데 PWM이 29.5%까지만 가는 이유는?", "150 mm/s 목표, Kp=50이면 P=7.5%p이고 feed-forward 22%를 더해 29.5%다."),
        ("10", "모터는 도는데 속도가 0이면 어디부터 보나?", "PB6/PB7 신호, TIM4 CNT, Encoder_Init/Update 순으로 본다."),
        ("11", "거리 0.19 m에서 STOP된 뒤 0.25 m가 되면 바로 풀리나?", "아니다. 0.30 m 이상 정상 샘플이 3회 필요하다."),
        ("12", "센서가 정상이어도 -25% 명령이 움직이지 않는 이유는?", "후방 센서가 없어 reverse_unprotected를 STOP으로 만드는 현재 안전 정책 때문이다."),
    ]
    for num, q, a in questions:
        heading(doc, 3, f"문제 {num}")
        add_body(doc, q)
        add_callout(doc, "해설", a, "tip")

    heading(doc, 1, "부록 D · 실제 보드 시험 체크리스트", page_break_before=True)
    for phase, items in [
        (
            "전원 투입 전",
            [
                "바퀴가 공중에 있고 차체가 고정되었는가",
                "모터 전원 극성·퓨즈·케이블 굵기를 확인했는가",
                "STM32와 BTS7960의 GND가 공통인가",
                "PA0/PA3/PE2/PE3/PB6/PB7/PA5/PB3 배선이 최신 코드와 일치하는가",
                "B1 USER 버튼을 누르지 않도록 표시했는가",
                "PB3 SWO 연결 가능성이 있는 SB9 실물 상태를 확인했는가",
                "초음파 ECHO가 STM32 입력에서 3.3 V 이하인가",
            ],
        ),
        (
            "통신 확인",
            [
                "self-test 성공",
                "COM 포트 확인",
                "echo 성공",
                "telemetry 수신",
            ],
        ),
        (
            "엔코더 확인",
            [
                "정방향 count 증가",
                "역방향 count 감소",
                "정지 중 count 안정",
                "여러 바퀴 count/rev 반복성 확보",
            ],
        ),
        (
            "초음파·안전 확인",
            [
                "정상 거리에서 RANGE_LOST가 사라지고 CLEAR 조건이 되는가",
                "0.20 m 미만에서 PWM=0·Enable=LOW·SAFE_STOP인가",
                "0.30 m 이상 정상 샘플 3회 전에는 STOP이 유지되는가",
                "sensor timeout/stale에서 모터가 정지하는가",
                "0.20 m E-STOP reset 절차와 안전 조건을 별도 검증했는가",
                "후진 명령은 현재 의도대로 SAFE_STOP인가",
            ],
        ),
        (
            "모터 확인",
            [
                "25% 0.5초 정방향 후 자동 정지",
                "음수 명령은 회전하지 않고 SAFE_STOP",
                "반대 PWM 채널이 동시에 활성화되지 않음",
                "Enable이 Disabled에서 LOW",
            ],
        ),
        (
            "PID 확인",
            [
                "목표 부호와 측정 부호 일치",
                "부하 증가 시 PWM 증가",
                "시험 종료 후 READY와 PWM 0",
                "온도·소음·진동 기록",
            ],
        ),
    ]:
        heading(doc, 2, phase)
        for item in items:
            add_bullet(doc, f"□ {item}")

    heading(doc, 1, "부록 E · 근거 파일과 참고자료", page_break_before=True)
    heading(doc, 2, "프로젝트 내부 기준")
    internal_sources = [
        ("ssacurity-stm32-drive.ioc", "핀, Clock Tree, TIM2/3/4/5, USART1, DMA, NVIC, FreeRTOS 설정"),
        ("Core/Src/main.c", "초기화 순서, SystemClock, TIM/UART/GPIO, 초음파 capture callback"),
        ("Core/Src/stm32f4xx_hal_msp.c", "GPIO AF, peripheral clock, DMA, IRQ priority"),
        ("Core/Src/stm32f4xx_it.c", "USART1, DMA2, TIM2, TIM6 interrupt handler"),
        ("Core/Src/stm32f4xx_hal_timebase_tim.c", "TIM6 기반 HAL 1 ms tick"),
        ("Core/Inc/FreeRTOSConfig.h", "1 kHz RTOS tick, interrupt priority 규칙"),
        ("Drivers/BSP/Src/motor.c", "BTS7960 PWM/Enable 드라이버"),
        ("Drivers/BSP/Src/encoder.c", "TIM4 count, wrap, 속도 환산"),
        ("Drivers/BSP/Src/ultrasonic.c", "TIM2 input capture, 10 µs TRIG, semaphore"),
        ("App/Src/pid.c", "PID와 anti-windup"),
        ("App/Src/task_ultrasonic.c", "60 ms 측정, 거리 환산, 3개 중앙값 필터"),
        ("App/Src/safety.c", "거리 임계값, 히스테리시스, 래치, reverse 차단"),
        ("App/Src/task_safety.c", "10 ms 안전 입력 수집과 Control/Motor 차단 요청"),
        ("App/Src/task_control.c", "10 ms 제어 루프, mode, fault, 50 ms safety gate"),
        ("App/Src/app_freertos.c", "4개 task 생성, 명령 연결, safety telemetry, 자동 정지"),
        ("Core/Src/comm_service.c", "진단 범위 검사, 명령 timeout, response-before-actuation"),
        ("docs/uart_protocol_v1.md", "메시지 ID와 payload 계약"),
        ("docs/handover_2026-07-28.md", "벤치 시험값; 단 LPWM 핀 설명은 현재 코드보다 오래됨"),
    ]
    for path, role in internal_sources:
        add_body(doc, f"{path} — {role}", bold_lead=path)

    heading(doc, 2, "공식 외부 자료")
    p = doc.add_paragraph()
    p.add_run("STMicroelectronics UM1670 — ")
    add_hyperlink(
        p,
        "Discovery kit with STM32F429ZI MCU user manual",
        "https://www.st.com/resource/en/user_manual/um1670-discovery-kit-with-stm32f429zi-mcu-stmicroelectronics.pdf",
    )
    p2 = doc.add_paragraph()
    p2.add_run("STMicroelectronics MB1075 E01 — ")
    add_hyperlink(
        p2,
        "STM32F429I-DISC1 board schematic",
        "https://www.st.com/resource/en/schematic_pack/mb1075-f429i-e01_schematic.pdf",
    )
    p3 = doc.add_paragraph()
    p3.add_run("Infineon — ")
    add_hyperlink(
        p3,
        "BTS7960 data sheet",
        "https://www.infineon.com/assets/row/public/documents/10/57/infineon-bts7960-ds-en.pdf",
    )

    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(7)
    closing.paragraph_format.space_after = Pt(0)
    lead = closing.add_run("마지막 한 문장  ")
    set_run_font(lead, size=8.5, color=GREEN, bold=True)
    sentence = closing.add_run(
        "모터 제어는 ‘안전하게 출력하고, 정확히 측정하고, 일정한 시간마다 둘의 차이를 줄이는 일’이다."
    )
    set_run_font(sentence, size=8.5, color=INK)

    # Core properties.
    props = doc.core_properties
    props.title = "완전 쌩뉴비를 위한 STM32 모터 제어 코드 한 권으로 끝내기"
    props.subject = "ssacurity-stm32-drive 프로젝트 모터 제어 입문서"
    props.author = "OpenAI Codex"
    props.keywords = "STM32, motor, BTS7960, PWM, encoder, PID, FreeRTOS, beginner"
    props.comments = "Generated from main/e8acfdc plus the current uncommitted safety and ultrasonic worktree on 2026-07-28."

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
